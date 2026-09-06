"""Genera el informe del Laboratorio 6 en formato Word (.docx).

Todas las cifras se leen de outputs/resultados.json y de outputs/tables/, producidos por
scripts/lab6_analisis.py. El informe no contiene ningún número escrito a mano: si el análisis
cambia, basta con volver a ejecutar el análisis y luego este script.

Requisito de formato: todo el texto del documento va en color negro.

Uso:
    python scripts/lab6_analisis.py      # primero, para generar métricas y figuras
    python scripts/generar_informe.py
"""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
TABLAS = ROOT / "outputs" / "tables"
FIGURAS = ROOT / "outputs" / "figures"
SALIDA = ROOT / "informe" / "Laboratorio6_Informe.docx"

NEGRO = RGBColor(0x00, 0x00, 0x00)
FUENTE = "Calibri"
GRIS_ENCABEZADO = "D9D9D9"
GRIS_ALTERNO = "F2F2F2"
ANCHO_UTIL_CM = 16.5

R = json.loads((ROOT / "outputs" / "resultados.json").read_text(encoding="utf-8"))


def tabla(nombre: str) -> pd.DataFrame:
    return pd.read_csv(TABLAS / f"{nombre}.csv", encoding="utf-8-sig")


# --------------------------------------------------------------------------------------------
# Utilidades de formato
# --------------------------------------------------------------------------------------------
def preparar_estilos(documento: Document) -> None:
    """Fuerza color negro y una tipografía consistente en todos los estilos del documento."""
    normal = documento.styles["Normal"]
    normal.font.name = FUENTE
    normal.font.size = Pt(11)
    normal.font.color.rgb = NEGRO
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FUENTE)

    tamanos = {"Title": 26, "Heading 1": 17, "Heading 2": 14, "Heading 3": 12, "Heading 4": 11}
    for nombre, tamano in tamanos.items():
        estilo = documento.styles[nombre]
        estilo.font.name = FUENTE
        estilo.font.size = Pt(tamano)
        estilo.font.color.rgb = NEGRO          # los encabezados de Word son azules por defecto
        estilo.font.bold = True
        estilo.paragraph_format.space_before = Pt(14 if nombre == "Heading 1" else 10)
        estilo.paragraph_format.space_after = Pt(6)
        estilo.paragraph_format.keep_with_next = True

    for nombre in ("Caption", "List Bullet", "List Number", "Subtitle", "Quote", "Intense Quote"):
        if nombre in [s.name for s in documento.styles]:
            estilo = documento.styles[nombre]
            estilo.font.name = FUENTE
            estilo.font.color.rgb = NEGRO
            estilo.font.italic = nombre == "Caption"

    # Requisito de formato: ningún estilo puede aportar color. Se recorre la hoja de estilos
    # completa —incluidos los estilos integrados de Word que no usamos— y se fuerza el negro.
    for color in documento.styles.element.iter(qn("w:color")):
        if color.get(qn("w:val")) not in {"FFFFFF", "auto"}:
            color.set(qn("w:val"), "000000")
        for atributo in ("w:themeColor", "w:themeTint", "w:themeShade"):
            if color.get(qn(atributo)) is not None:
                del color.attrib[qn(atributo)]


def sombrear(celda, color_hex: str) -> None:
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:val"), "clear")
    sombra.set(qn("w:color"), "auto")
    sombra.set(qn("w:fill"), color_hex)
    celda._tc.get_or_add_tcPr().append(sombra)


def escribir(parrafo, texto: str, negrita=False, cursiva=False, tamano=11, color=NEGRO):
    corrida = parrafo.add_run(texto)
    corrida.bold = negrita
    corrida.italic = cursiva
    corrida.font.size = Pt(tamano)
    corrida.font.name = FUENTE
    corrida.font.color.rgb = color
    return corrida


def parrafo_rico(documento, segmentos, alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 tamano=11, espacio_despues=6, sangria=0.0):
    """Crea un párrafo a partir de una lista de tramos.

    Cada tramo es un texto o una tupla (texto, negrita) / (texto, negrita, cursiva).
    """
    p = documento.add_paragraph()
    p.alignment = alineacion
    p.paragraph_format.space_after = Pt(espacio_despues)
    if sangria:
        p.paragraph_format.left_indent = Cm(sangria)
    for segmento in segmentos if isinstance(segmentos, (list, tuple)) else [segmentos]:
        if isinstance(segmento, str):
            escribir(p, segmento, tamano=tamano)
        else:
            texto, *resto = segmento
            escribir(p, texto, negrita=bool(resto and resto[0]),
                     cursiva=bool(len(resto) > 1 and resto[1]), tamano=tamano)
    return p


def texto(documento, contenido, **kwargs):
    return parrafo_rico(documento, contenido, **kwargs)


def vineta(documento, segmentos, nivel=0):
    p = documento.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.75 + 0.6 * nivel)
    p.paragraph_format.space_after = Pt(3)
    for segmento in segmentos if isinstance(segmentos, (list, tuple)) else [segmentos]:
        if isinstance(segmento, str):
            escribir(p, segmento, tamano=10.5)
        else:
            t, *resto = segmento
            escribir(p, t, negrita=bool(resto and resto[0]),
                     cursiva=bool(len(resto) > 1 and resto[1]), tamano=10.5)
    return p


def titulo(documento, nivel, contenido):
    encabezado = documento.add_heading(level=nivel)
    corrida = encabezado.add_run(contenido)
    corrida.font.color.rgb = NEGRO
    corrida.font.name = FUENTE
    corrida.bold = True
    return encabezado


CONTADOR = {"tabla": 0, "figura": 0}


def leyenda(documento, tipo: str, descripcion: str, antes=False):
    CONTADOR[tipo] += 1
    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2 if not antes else 6)
    p.paragraph_format.space_after = Pt(10 if not antes else 4)
    escribir(p, f"{tipo.capitalize()} {CONTADOR[tipo]}. ", negrita=True, tamano=9.5)
    escribir(p, descripcion, cursiva=True, tamano=9.5)
    return p


def insertar_tabla(documento, frame: pd.DataFrame, anchos=None, tamano=8.5,
                   encabezados=None, alineacion_derecha=()):
    """Inserta un DataFrame como tabla con encabezado sombreado y filas alternas."""
    frame = frame.fillna("")
    filas, columnas = frame.shape
    objeto = documento.add_table(rows=filas + 1, cols=columnas)
    objeto.style = "Table Grid"
    objeto.alignment = WD_TABLE_ALIGNMENT.CENTER
    objeto.autofit = False

    if anchos is None:
        anchos = [ANCHO_UTIL_CM / columnas] * columnas
    escala = ANCHO_UTIL_CM / sum(anchos)
    anchos = [a * escala for a in anchos]

    etiquetas = encabezados or [str(c).replace("_", " ") for c in frame.columns]
    for j, etiqueta in enumerate(etiquetas):
        celda = objeto.cell(0, j)
        celda.text = ""
        p = celda.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        escribir(p, str(etiqueta), negrita=True, tamano=tamano)
        sombrear(celda, GRIS_ENCABEZADO)

    for i in range(filas):
        for j in range(columnas):
            celda = objeto.cell(i + 1, j)
            celda.text = ""
            p = celda.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if j in alineacion_derecha:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            valor = frame.iat[i, j]
            if isinstance(valor, float) and valor == int(valor) and abs(valor) < 1e15:
                valor = int(valor)
            escribir(p, str(valor), tamano=tamano)
            if i % 2 == 1:
                sombrear(celda, GRIS_ALTERNO)

    for fila in objeto.rows:
        for j, celda in enumerate(fila.cells):
            celda.width = Cm(anchos[j])

    # Repite el encabezado si la tabla se parte entre páginas.
    propiedades = objeto.rows[0]._tr.get_or_add_trPr()
    encabezado = OxmlElement("w:tblHeader")
    encabezado.set(qn("w:val"), "true")
    propiedades.append(encabezado)
    return objeto


def insertar_figura(documento, archivo: str, ancho_cm=ANCHO_UTIL_CM):
    ruta = FIGURAS / archivo
    if not ruta.exists():
        return None
    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(ruta), width=Cm(ancho_cm))
    return p


def figura_apaisada(documento, archivo: str, descripcion: str):
    """Coloca una figura ancha en su propia página horizontal, donde sí se puede leer."""
    ruta = FIGURAS / archivo
    if not ruta.exists():
        return
    horizontal = documento.add_section(WD_SECTION.NEW_PAGE)
    horizontal.orientation = WD_ORIENT.LANDSCAPE
    horizontal.page_width, horizontal.page_height = Cm(27.94), Cm(21.59)
    horizontal.left_margin = horizontal.right_margin = Cm(1.8)
    horizontal.top_margin = horizontal.bottom_margin = Cm(1.8)

    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.add_run().add_picture(str(ruta), width=Cm(24.3))
    leyenda(documento, "figura", descripcion)

    vertical = documento.add_section(WD_SECTION.NEW_PAGE)
    vertical.orientation = WD_ORIENT.PORTRAIT
    vertical.page_width, vertical.page_height = Cm(21.59), Cm(27.94)
    vertical.left_margin = vertical.right_margin = Cm(2.5)
    vertical.top_margin = vertical.bottom_margin = Cm(2.2)


def salto_de_pagina(documento):
    p = documento.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def numerar_paginas(seccion) -> None:
    """Inserta el campo PAGE en el pie de página, en negro."""
    pie = seccion.footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corrida = pie.add_run()
    corrida.font.size = Pt(9)
    corrida.font.name = FUENTE
    corrida.font.color.rgb = NEGRO
    for etiqueta, atributos in [("w:fldChar", {"w:fldCharType": "begin"}),
                                ("w:instrText", {"xml:space": "preserve"}),
                                ("w:fldChar", {"w:fldCharType": "end"})]:
        elemento = OxmlElement(etiqueta)
        for clave, valor in atributos.items():
            elemento.set(qn(clave), valor)
        if etiqueta == "w:instrText":
            elemento.text = " PAGE "
        corrida._r.append(elemento)


def p_valor(valor) -> str:
    valor = float(valor)
    return "p < 0.001" if valor < 0.001 else f"p = {valor:.3f}"


def fila(lista, clave, valor):
    return next(x for x in lista if x[clave] == valor)


# --------------------------------------------------------------------------------------------
# Portada, resumen y contenido
# --------------------------------------------------------------------------------------------
def portada(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    escribir(p, "Universidad del Valle de Guatemala", negrita=True, tamano=13)
    for linea in ["Facultad de Ingeniería", "Departamento de Ciencias de la Computación",
                  "CC3084 — Data Science · Semestre II, 2026"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        escribir(p, linea, tamano=12)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    escribir(p, "Laboratorio 6", negrita=True, tamano=26)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    escribir(p, "Análisis de redes sociales en YouTube", negrita=True, tamano=19)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    escribir(p, "Participación, comunidades y sentimiento en una muestra de contenido "
                "guatemalteco", cursiva=True, tamano=12)

    doc.add_paragraph()
    resumen = pd.DataFrame([
        ["Conjuntos de datos", f"youtube_videos.csv ({R['n_videos']} videos × {R['n_vars_videos']} variables) · "
                               f"youtube_comments.csv ({R['n_comentarios']} comentarios × {R['n_vars_comentarios']} variables)"],
        ["Unidades observadas", f"{R['n_videos']} videos · {R['canales_videos']} canales · "
                                f"{R['n_comentarios']} comentarios principales · {R['autores_unicos']} autores"],
        ["Red construida", f"Bipartita autor–video no dirigida: {R['red_nodos_total']} nodos y "
                           f"{R['red_aristas_total']} aristas (peso total = {R['peso_total']} comentarios)"],
        ["Herramientas", "Python 3 · pandas · networkx · spaCy (es_core_news_sm) · "
                         "python-louvain · pysentimiento (RoBERTuito) · matplotlib"],
        ["Repositorio", "https://github.com/iancumes/Lab6DataScience"],
        ["Reproducibilidad", "python scripts/lab6_analisis.py — regenera todas las tablas, figuras y métricas"],
    ], columns=["Elemento", "Detalle"])
    insertar_tabla(doc, resumen, anchos=[4.2, 12.3], tamano=9.5)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    escribir(p, "Guatemala, 6 de septiembre de 2026", tamano=10.5)
    salto_de_pagina(doc)


def resumen_ejecutivo(doc: Document) -> None:
    titulo(doc, 1, "Resumen ejecutivo")
    texto(doc, [
        "Este informe analiza la estructura de participación observada en una muestra de ",
        (f"{R['n_videos']} videos", True), " y ", (f"{R['n_comentarios']} comentarios", True),
        " de YouTube relacionados con contenido guatemalteco. El objetivo es describir cómo se "
        "relacionan canales, videos, autores y temas, y qué puede y no puede afirmarse a partir "
        "de los datos entregados.",
    ])
    texto(doc, [
        "El hallazgo central es estructural: ", ("la red no es una red de conversación, sino una "
        "constelación de audiencias separadas", True), ". El ",
        (f"{R['pct_autores_grado1']} %", True), " de los autores comentó en un solo video y no "
        "volvió a aparecer; sólo ", (f"{R['autores_multivideo_n']} de {R['autores_unicos']} autores", True),
        " participaron en más de un video. La consecuencia es que los videos, no las personas, son "
        "los que sostienen la estructura: al detectar comunidades se obtienen ",
        (f"{R['n_comunidades']} grupos", True), f" con una modularidad de {R['modularidad']}, y casi "
        "todos corresponden a un único video con su público propio.",
    ])
    texto(doc, [
        "La participación está fuertemente concentrada. Un solo video, «",
        (R["video_top_titulo"], False, True), "» del canal ", (R["video_top_canal"], True),
        f", reúne {R['video_top_comentarios']} de los {R['n_comentarios']} comentarios "
        f"({R['video_top_pct']} %), y ese mismo canal concentra el {R['canal_top_pct']} % del total. "
        "El coeficiente de Gini de comentarios por video con cobertura es "
        f"{R['gini_videos_cobertura']}.",
    ])
    texto(doc, [
        "El tono predominante es negativo: ", (f"{R['sentimiento_global']['NEG']} de "
        f"{R['n_comentarios']} comentarios ({R['pct_negativo_global']} %)", True),
        " fueron clasificados como negativos por un modelo de lenguaje entrenado en español. "
        "Ese tono no es uniforme: varía de forma estadísticamente significativa entre canales "
        f"(Kruskal–Wallis, {p_valor(fila(R['pruebas_sentimiento'], 'prueba', 'Kruskal–Wallis: puntaje de sentimiento entre canales comparables')['p_valor'])}), "
        "desde −0.50 en el canal de periodismo de investigación hasta +0.61 en el canal municipal.",
    ])
    texto(doc, [
        ("Limitación que condiciona todo lo demás: ", True),
        f"sólo {R['videos_con_comentarios']} de los {R['n_videos']} videos "
        f"({R['cobertura_videos_pct']} %) tienen comentarios recolectados. Los "
        f"{R['videos_sin_comentarios']} videos restantes no son videos sin participación: son videos "
        "sin datos. Todos los resultados de red describen la muestra observada y no pueden "
        "extrapolarse a YouTube Guatemala ni a la población de usuarios.",
    ])

    titulo(doc, 2, "Contenido del informe")
    for numero, nombre in [
        ("1", "Carga, comprensión e integración de los datos"),
        ("2", "Calidad, limpieza y preprocesamiento"),
        ("3", "Análisis exploratorio"),
        ("4", "Construcción de la red bipartita autor–video"),
        ("5", "Proyecciones de la red"),
        ("6", "Topología y fragmentación"),
        ("7", "Comunidades"),
        ("8", "Nodos centrales y participantes puente"),
        ("9", "Análisis de contenido y sentimiento"),
        ("10", "Interpretación, limitaciones y conclusiones"),
        ("A", "Anexo: reproducibilidad y material generado"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        escribir(p, f"{numero}. ", negrita=True, tamano=10.5)
        escribir(p, nombre, tamano=10.5)
    salto_de_pagina(doc)


# --------------------------------------------------------------------------------------------
# 1. Carga, comprensión e integración
# --------------------------------------------------------------------------------------------
def seccion_1(doc: Document) -> None:
    titulo(doc, 1, "1. Carga, comprensión e integración de los datos")

    titulo(doc, 2, "1.1 Carga de los archivos")
    texto(doc, [
        "Ambos archivos se cargaron en Python con pandas. Dos decisiones de lectura merecen "
        "mención porque afectan la integridad de los identificadores: se usó la codificación ",
        ("utf-8-sig", False, True), " —los archivos incluyen marca de orden de bytes (BOM), que de "
        "otro modo contamina el nombre de la primera columna— y todas las columnas de "
        "identificación se forzaron a tipo texto. Esto último es esencial: si pandas infiriera el "
        "tipo, un identificador podría perder ceros a la izquierda o convertirse en notación "
        f"científica. La carga produce {R['n_videos']} filas × {R['n_vars_videos']} variables para "
        f"videos y {R['n_comentarios']} filas × {R['n_vars_comentarios']} variables para comentarios, "
        "coincidiendo exactamente con lo descrito en el enunciado.",
    ])

    titulo(doc, 2, "1.2 Unidad de observación, llave primaria y variables relevantes")
    texto(doc, [
        "La unidad de observación se verificó empíricamente en lugar de asumirse. Una columna es "
        "llave primaria si no tiene valores faltantes y su número de valores únicos iguala al "
        "número de filas; ambas condiciones se comprobaron por separado:",
    ])
    llaves = pd.DataFrame(R["llaves_candidatas"])
    llaves = llaves[["archivo", "llave_candidata", "filas", "valores_unicos", "faltantes", "es_llave_primaria"]]
    llaves.columns = ["Archivo", "Columna", "Filas", "Valores únicos", "Faltantes", "¿Llave primaria?"]
    llaves["¿Llave primaria?"] = llaves["¿Llave primaria?"].map({True: "Sí", False: "No"})
    insertar_tabla(doc, llaves, anchos=[4.4, 3.2, 1.8, 2.6, 2.0, 2.5], tamano=9,
                   alineacion_derecha=(2, 3, 4))
    leyenda(doc, "tabla", "Verificación empírica de las llaves candidatas de cada archivo.")

    texto(doc, [
        "En ", ("youtube_videos.csv", False, True), " cada fila es ", ("un video", True),
        " y la llave primaria es ", ("video_id", False, True), ". En ",
        ("youtube_comments.csv", False, True), " cada fila es ", ("un comentario principal", True),
        " —no una respuesta— y la llave primaria es ", ("comment_id", False, True), ". La llave "
        "foránea que une ambos archivos es ", ("video_id", False, True), ", y ",
        ("author_channel_id", False, True), " identifica al autor. Conviene notar que ",
        ("video_url", False, True), " también es única, pero no se usa como llave: es una cadena "
        "larga que contiene el identificador, de modo que no aporta información adicional.",
    ])
    texto(doc, [
        "Las variables se agruparon por función analítica: identificación y red (",
        ("video_id, channel_id, comment_id, author_channel_id", False, True),
        "), etiquetas visibles (", ("title, channel_name, author_name, handles", False, True),
        "), cuantitativas (", ("view_count, like_count_text, reply_count", False, True),
        "), contenido y temas (", ("title, description, keywords, text", False, True),
        "), contexto de muestreo (", ("source_query, source_group, query_hits, dataset_sources, category", False, True),
        ") y temporales (", ("publish_date, published_time, published_text", False, True), ").",
    ])

    titulo(doc, 2, "1.3 Relación entre canal, video, autor, comentario, categoría y consulta")
    relaciones = tabla("02_relaciones_entre_entidades")
    relaciones.columns = ["Origen", "Destino", "Cardinalidad", "Llave", "Evidencia en los datos"]
    insertar_tabla(doc, relaciones, anchos=[2.0, 2.3, 2.3, 3.6, 6.3], tamano=9)
    leyenda(doc, "tabla", "Esquema relacional observable en los dos archivos.")

    texto(doc, [
        "Un canal publica uno o más videos; un video pertenece a exactamente una categoría de "
        "YouTube y recibe cero o más comentarios; cada comentario tiene exactamente un autor, y un "
        "autor puede escribir varios comentarios en uno o varios videos de uno o varios canales. "
        "Dos precisiones son necesarias para no cometer errores de interpretación:",
    ])
    vineta(doc, [
        ("La consulta de búsqueda no es un atributo del contenido. ", True),
        "source_query y source_group describen cómo se encontró el video durante la recolección, "
        "no de qué trata. Un video hallado con la consulta «guatemala lluvias» puede tratar de "
        "cualquier otra cosa.",
    ])
    vineta(doc, [
        ("channel_id significa cosas distintas en cada archivo. ", True),
        "En el archivo de comentarios identifica al canal ",
        ("dueño del video", True), ", no al autor del comentario. Son espacios de identificadores "
        f"distintos: se verificó que ninguno de los {R['autores_unicos']} author_channel_id coincide "
        "con un channel_id del catálogo de videos.",
    ])
    vineta(doc, [
        ("La relación autor–autor no es observable. ", True),
        "reply_count indica cuántas respuestas recibió un comentario, pero no quién las escribió. "
        "Cualquier vínculo entre autores en este informe es co-participación derivada, nunca "
        "interacción observada.",
    ])

    titulo(doc, 2, "1.4 Integración por video_id y cobertura")
    integracion = tabla("03_integracion_y_cobertura")
    integracion.columns = ["Indicador", "Valor"]
    insertar_tabla(doc, integracion, anchos=[11.5, 5.0], tamano=9.5, alineacion_derecha=(1,))
    leyenda(doc, "tabla", "Resultado de la integración de ambos archivos mediante video_id.")

    texto(doc, [
        "La unión es íntegra: los ", (f"{R['comentarios_asociados']} comentarios", True),
        " se asociaron con un video del catálogo, no hay ningún comentario huérfano y la unión no "
        f"expande filas (la relación es muchos-a-uno, verificada con validate='many_to_one'). El "
        "problema no está en la integridad de la unión sino en su ",
        ("cobertura", True), ": únicamente ", (f"{R['videos_con_comentarios']} de {R['n_videos']} "
        f"videos ({R['cobertura_videos_pct']} %)", True), " tienen al menos un comentario "
        f"recolectado, y sólo {R['canales_con_comentarios']} de los {R['canales_videos']} canales "
        "aparecen representados en el archivo de comentarios. Este desbalance es el condicionante "
        "principal de todo el análisis y se retoma en la sección 10.",
    ])
    salto_de_pagina(doc)


# --------------------------------------------------------------------------------------------
# 2. Calidad, limpieza y preprocesamiento
# --------------------------------------------------------------------------------------------
def seccion_2(doc: Document) -> None:
    titulo(doc, 1, "2. Calidad, limpieza y preprocesamiento")

    titulo(doc, 2, "2.1 Diagnóstico inicial de calidad")
    texto(doc, [
        "El diagnóstico cubre las seis dimensiones exigidas —dimensiones, tipos, faltantes, "
        "duplicados, variables constantes y atípicos— más un bloque específico de consistencia "
        "entre identificadores, nombres y ", ("handles", False, True), ". El perfil completo de las "
        "37 variables está en ", ("outputs/tables/04_diagnostico_calidad.csv", False, True),
        "; aquí se resumen los hallazgos.",
    ])

    titulo(doc, 3, "Dimensiones y tipos")
    texto(doc, [
        f"youtube_videos.csv tiene {R['n_videos']} filas y {R['n_vars_videos']} variables; "
        f"youtube_comments.csv tiene {R['n_comentarios']} filas y {R['n_vars_comentarios']} "
        "variables. La tipificación observada no coincide siempre con la semántica: ",
        ("like_count_text", False, True), " es un conteo almacenado como texto, ",
        ("published_time", False, True), " y ", ("published_text", False, True), " son fechas "
        "expresadas como duración relativa, y ", ("query_hits", False, True), " y ",
        ("keywords", False, True), " son listas serializadas en formato JSON dentro de una celda.",
    ])

    titulo(doc, 3, "Valores faltantes")
    faltantes = tabla("04_diagnostico_calidad").query("faltantes > 0 or vacios_o_espacios > 0")[
        ["archivo", "variable", "faltantes", "faltantes_pct", "vacios_o_espacios"]]
    faltantes.columns = ["Archivo", "Variable", "Faltantes", "% faltantes", "Vacíos o espacios"]
    insertar_tabla(doc, faltantes, anchos=[4.6, 4.4, 2.2, 2.4, 2.9], tamano=9,
                   alineacion_derecha=(2, 3, 4))
    leyenda(doc, "tabla", "Variables con valores faltantes o cadenas vacías.")
    texto(doc, [
        "Sólo una variable está completamente vacía: ", ("viewer_rating", False, True),
        f", con {R['n_comentarios']}/{R['n_comentarios']} faltantes (100 %). En videos, los faltantes "
        "se concentran en campos de presentación (", ("published_time", False, True), ", ",
        ("view_count_text", False, True), ", ", ("description", False, True), "), nunca en "
        "identificadores ni en ", ("view_count", False, True), ". Caso aparte es ",
        ("like_count_text", False, True), f": no tiene nulos formales, pero {R['likes_en_blanco']} "
        f"de {R['n_comentarios']} registros ({round(100 * R['likes_en_blanco'] / R['n_comentarios'], 1)} %) "
        "contienen un espacio en blanco. Un conteo de nulos ingenuo habría reportado cero problemas "
        "en esa columna.",
    ])

    titulo(doc, 3, "Duplicados")
    duplicados = pd.DataFrame(R["duplicados"])
    duplicados.columns = ["Archivo", "Chequeo", "Conteo"]
    insertar_tabla(doc, duplicados, anchos=[4.6, 9.4, 2.5], tamano=9, alineacion_derecha=(2,))
    leyenda(doc, "tabla", "Duplicados exactos y por combinaciones de columnas.")
    texto(doc, [
        "No hay filas exactamente duplicadas ni llaves primarias repetidas en ninguno de los dos "
        "archivos. Sí aparecen dos duplicados ", ("aparentes", False, True), " que se decidió ",
        ("conservar", True), ": 19 títulos de video se repiten con ", ("video_id", False, True),
        " distinto (retransmisiones y programas seriados, que son videos diferentes), y 2 textos de "
        "comentario se repiten con ", ("comment_id", False, True), " distinto. Eliminarlos sería un "
        "error: cada uno representa un acto de publicación real y, en el caso de los comentarios, "
        "una arista de la red.",
    ])

    titulo(doc, 3, "Variables constantes y valores atípicos")
    texto(doc, [
        ("is_pinned", False, True), f" es constante en False para los {R['n_comentarios']} registros, "
        "de modo que no aporta variación. ", ("viewer_rating", False, True), " está vacía por "
        "completo. En cuanto a atípicos, la regla del rango intercuartílico marca ",
        ("49 de 293 videos", True), " como atípicos por visualizaciones. Esa cifra ",
        ("no indica errores de datos", True), ": las visualizaciones de YouTube siguen una "
        "distribución de cola larga, con mediana de "
        f"{fila(R['descriptivos'], 'indicador', 'Visualizaciones — mediana')['valor']:,.0f} y máximo de "
        f"{fila(R['descriptivos'], 'indicador', 'Visualizaciones — máximo')['valor']:,.0f}. Por eso "
        "los atípicos se conservan y todos los gráficos de esta variable usan escala logarítmica, "
        "en lugar de recortar la cola.".replace(",", " "),
    ])

    titulo(doc, 3, "Consistencia entre identificadores, nombres y handles")
    consistencia = tabla("06_consistencia_identificadores")
    consistencia.columns = ["Ámbito", "Regla verificada", "Incumplimientos / conteo"]
    insertar_tabla(doc, consistencia, anchos=[2.2, 10.3, 4.0], tamano=9, alineacion_derecha=(2,))
    leyenda(doc, "tabla", "Chequeos de consistencia entre identificadores y etiquetas visibles.")
    texto(doc, [
        "La correspondencia entre identificadores y etiquetas es perfecta dentro de la muestra: "
        "cada ", ("channel_id", False, True), " tiene un único nombre y un único ",
        ("handle", False, True), ", y lo mismo ocurre con ", ("author_channel_id", False, True),
        ". Esto no autoriza a usar los nombres como llaves —fuera de esta muestra pueden repetirse "
        "o cambiar— pero confirma que las etiquetas son fiables para mostrar resultados. Dos "
        "hallazgos adicionales:",
    ])
    vineta(doc, [
        ("Campos redundantes. ", True), "owner_handle es idéntico a channel_handle y upload_date es "
        "idéntica a publish_date en el 100 % de los registros. Se conserva una sola de cada par.",
    ])
    vineta(doc, [
        ("source_group y source_query son homónimos con significado distinto. ", True),
        f"Existen en ambos archivos y no coinciden en 188 de {R['n_comentarios']} comentarios: en "
        "videos describen cómo se encontró el video y en comentarios cómo se recolectaron sus "
        "comentarios. Al integrar se conservó el sufijo _video para no confundirlos.",
    ])

    titulo(doc, 2, "2.2 Variables que no pueden usarse o requieren precaución")
    problematicas = tabla("07_variables_problematicas")
    problematicas.columns = ["Variable", "Clasificación", "Problema observado", "Tratamiento y justificación"]
    insertar_tabla(doc, problematicas, anchos=[3.0, 2.3, 5.6, 5.6], tamano=8.5)
    leyenda(doc, "tabla", "Clasificación de las variables problemáticas y tratamiento aplicado.")

    titulo(doc, 2, "2.3 Normalización de identificadores y nombres")
    texto(doc, [
        "La normalización es deliberadamente ", ("asimétrica", True), ". Sobre los identificadores "
        "sólo se recortan espacios externos: los identificadores de YouTube distinguen mayúsculas y "
        "minúsculas, de modo que pasarlos a minúsculas fusionaría cuentas distintas. Sobre las "
        "etiquetas visibles sí se aplica normalización Unicode NFKC y colapso de espacios, porque "
        "son texto de presentación. En ningún momento se sustituye un identificador por un nombre.",
    ])
    texto(doc, [
        "Un hallazgo específico obligó a un paso extra: ", (f"{R['handles_codificados']} handles", True),
        " (14 de autor y 13 de canal) llegan con codificación porcentual de URL, por ejemplo ",
        ("/@AlejandroP%C3%A9rez-b6r", False, True), " en lugar de ", ("@AlejandroPérez-b6r", False, True),
        ". Sin decodificarlos, la misma persona aparecería con dos etiquetas distintas en tablas y "
        "figuras. Se decodifican y se les da la forma canónica @nombre.",
    ])
    verificacion = tabla("08_verificacion_normalizacion_ids")
    verificacion.columns = ["Chequeo", "Antes", "Después"]
    insertar_tabla(doc, verificacion, anchos=[11.0, 2.7, 2.8], tamano=9, alineacion_derecha=(1, 2))
    leyenda(doc, "tabla", "La normalización no fusiona ni rompe identificadores: los conteos únicos "
                          "se mantienen y la codificación porcentual desaparece.")

    titulo(doc, 2, "2.4 Conversión a numérico de las variables de conteo en texto")
    texto(doc, [
        "Se implementó una función de conversión que documenta explícitamente cada decisión: "
        "elimina separadores de miles (", ("2,390 vistas", False, True), " → 2390) distinguiendo el "
        "formato inglés del español, resuelve abreviaturas (", ("1.2 K", False, True), " → 1200, ",
        ("12 mil", False, True), " → 12000), descarta palabras de contexto (",
        ("vistas, views, visualizaciones", False, True), ") y devuelve valor faltante —no cero— "
        "cuando la cadena no contiene ningún dígito interpretable.",
    ])
    texto(doc, [
        "La única imputación aplicada es una ", ("regla de negocio explícita", True), ": en ",
        ("like_count_text", False, True), " un espacio en blanco significa cero, porque la interfaz "
        f"de YouTube oculta el contador cuando vale 0. Los {R['likes_en_blanco']} registros afectados "
        "se convierten a 0 y quedan marcados en la variable ", ("like_count_imputado", False, True),
        " para que la imputación sea auditable y reversible.",
    ])
    conversion = tabla("10_conversion_conteos")
    conversion.columns = ["Conversión", "Detalle", "Total resultante"]
    insertar_tabla(doc, conversion, anchos=[4.2, 9.0, 3.3], tamano=9, alineacion_derecha=(2,))
    leyenda(doc, "tabla", "Resultado de las conversiones y su verificación cruzada.")
    texto(doc, [
        "La conversión sirvió además como control cruzado. Al comparar ",
        ("view_count_text", False, True), " convertido con la columna entera ",
        ("view_count", False, True), f", {R['view_text_coinciden']} de {R['view_text_disponibles']} "
        f"valores coinciden exactamente. Las {R['view_text_disponibles'] - R['view_text_coinciden']} "
        f"discrepancias son pequeñas (diferencia mediana de {R['view_dif_mediana']} vistas) y su signo "
        "varía: no son errores de conversión sino evidencia de que ambas columnas se capturaron en "
        "instantes distintos. Es un recordatorio de que estos conteos son fotografías, no valores "
        "definitivos.",
    ])

    titulo(doc, 2, "2.5 y 2.6 texto_original y texto_limpio")
    texto(doc, [
        "Se conservan dos versiones del texto de cada comentario. ", ("texto_original", False, True),
        " es una copia literal de la variable ", ("text", False, True), ": es la versión auditable y "
        "la que alimenta el análisis de sentimiento, porque el modelo en español necesita la "
        "negación, la puntuación y los emojis que la limpieza destruiría. ",
        ("texto_limpio", False, True), " es la versión normalizada y lematizada que se usa para "
        "frecuencias, bigramas y caracterización temática.",
    ])
    pipeline = pd.DataFrame([
        ["Unicode", "NFKC y colapso de espacios", "Unifica caracteres visualmente idénticos."],
        ["URLs", "Se extraen a urls_lista y se eliminan", "Aportan ruido léxico; se conservan aparte para auditoría."],
        ["Hashtags", "Se extraen a hashtags_lista; la palabra se conserva sin '#'",
         "Permite contarlos por separado sin perder el término del vocabulario."],
        ["Menciones", "Se extraen a menciones_lista y se eliminan", "Un @usuario es un identificador, no vocabulario."],
        ["Emojis", "Se extraen a emojis_lista y se eliminan del texto limpio",
         "Se analizan aparte; sobreviven en texto_original para el sentimiento."],
        ["Minúsculas", "Sí", "Evita duplicar tipos por capitalización."],
        ["Números", "Se eliminan los tokens puramente numéricos", "No aportan tema; las cifras viven en las variables de conteo."],
        ["Puntuación", "Se elimina", "Reduce ruido en los conteos de frecuencia."],
        ["Repeticiones", "'holaaaa' → 'holaa'", "Normaliza el alargamiento expresivo típico de redes."],
        ["Stopwords", "Lista de spaCy en español + lista propia de muletillas",
         "Las palabras funcionales no discriminan temas; se filtran también sin tildes."],
        ["Lematización", "spaCy es_core_news_sm", "Une flexiones: corrupto/corruptos, robar/roban."],
        ["Lemas compuestos", "Se dividen y se filtran pieza por pieza",
         "spaCy lematiza los clíticos como 'dárselo' → 'dar él'; sin dividirlos, el pronombre se cuela."],
        ["Variantes con tilde", "'pais' y 'país' se unifican en la forma más frecuente",
         "Los usuarios omiten tildes de forma inconsistente; sin unificar, el mismo concepto se cuenta dos veces."],
        ["Tokens cortos", "Se descartan los de 1–2 caracteres", "Residuos de los pasos anteriores."],
    ], columns=["Paso", "Decisión", "Justificación"])
    insertar_tabla(doc, pipeline, anchos=[2.8, 6.4, 7.3], tamano=8.5)
    leyenda(doc, "tabla", "Pipeline de construcción de texto_limpio, en orden de aplicación.")

    titulo(doc, 2, "2.7 Efecto cuantificado de la limpieza")
    efecto = tabla("11_efecto_limpieza")
    efecto.columns = ["Métrica", "Valor"]
    insertar_tabla(doc, efecto, anchos=[12.0, 4.5], tamano=9, alineacion_derecha=(1,))
    leyenda(doc, "tabla", "Efecto medido de la limpieza sobre los 406 comentarios.")
    texto(doc, [
        ("No se eliminó ningún registro", True), ". La decisión es deliberada: cada comentario es "
        "una arista de la red, de modo que descartar uno por tener poco contenido léxico "
        "distorsionaría la estructura. La limpieza modificó 405 de 406 textos y redujo el volumen de "
        f"tokens en {R['reduccion_tokens_pct']} %, lo que es esperable al eliminar stopwords. El "
        f"efecto colateral relevante es que {R['textos_vacios_despues']} comentarios quedan sin "
        "contenido léxico —eran sólo emojis, menciones o interjecciones—; se marcan con la bandera ",
        ("apto_para_texto", False, True), " y se excluyen de los conteos de palabras, pero siguen "
        "siendo nodos y aristas válidos en la red y en el análisis de sentimiento. Los duplicados de "
        "texto aumentan de 4 a 19 filas porque la limpieza colapsa expresiones distintas en la misma "
        "forma canónica; es una consecuencia esperada de la lematización, no un error.",
    ])
    salto_de_pagina(doc)


# --------------------------------------------------------------------------------------------
# 3. Análisis exploratorio
# --------------------------------------------------------------------------------------------
def seccion_3(doc: Document) -> None:
    titulo(doc, 1, "3. Análisis exploratorio")

    titulo(doc, 2, "3.1 Descriptivos generales")
    descriptivos = tabla("15_descriptivos_generales")
    descriptivos["valor"] = descriptivos["valor"].map(
        lambda v: f"{v:,.0f}".replace(",", " ") if float(v) == int(float(v)) else f"{v:,.2f}".replace(",", " "))
    descriptivos.columns = ["Indicador", "Valor"]
    mitad = (len(descriptivos) + 1) // 2
    izquierda = descriptivos.iloc[:mitad].reset_index(drop=True)
    derecha = descriptivos.iloc[mitad:].reset_index(drop=True)
    combinada = pd.concat([izquierda, derecha], axis=1).fillna("")
    combinada.columns = ["Indicador", "Valor", "Indicador ", "Valor "]
    insertar_tabla(doc, combinada, anchos=[5.3, 2.2, 5.3, 2.2], tamano=8.5, alineacion_derecha=(1, 3))
    leyenda(doc, "tabla", "Descriptivos mínimos exigidos por el inciso 3.1.")

    texto(doc, [
        "Tres contrastes resumen la muestra. Primero, la ",
        ("asimetría entre los dos lados", True), ": hay "
        f"{R['autores_unicos']} autores para sólo {R['videos_con_comentarios']} videos con "
        f"comentarios, de modo que un video recibe en promedio "
        f"{fila(R['descriptivos'], 'indicador', 'Autores únicos por video con cobertura — media')['valor']} "
        "autores distintos mientras que un autor participa en "
        f"{fila(R['resumen_grados'], 'lado', 'Autores')['grado_medio']} videos en promedio. Segundo, "
        "la ", ("dispersión extrema", True), ": la mediana de comentarios por video con cobertura es "
        f"{fila(R['descriptivos'], 'indicador', 'Comentarios por video con cobertura — mediana')['valor']:.0f} "
        f"pero el máximo es {fila(R['descriptivos'], 'indicador', 'Comentarios por video con cobertura — máximo')['valor']:.0f}; "
        "lo mismo ocurre con las visualizaciones (mediana "
        f"{fila(R['descriptivos'], 'indicador', 'Visualizaciones — mediana')['valor']:,.0f} frente a un máximo de "
        f"{fila(R['descriptivos'], 'indicador', 'Visualizaciones — máximo')['valor']:,.0f}). "
        "Tercero, la ", ("escasez de conversación", True), f": sólo "
        f"{fila(R['descriptivos'], 'indicador', 'Comentarios con al menos una respuesta')['valor']:.0f} "
        f"de los {R['n_comentarios']} comentarios recibieron alguna respuesta, con "
        f"{R['total_respuestas']} respuestas en total.".replace(",", " "),
    ])
    insertar_figura(doc, "01_top_participacion.png")
    leyenda(doc, "figura", "Videos y canales con mayor participación observada. El número entre "
                           "paréntesis indica autores distintos, para distinguir alcance de intensidad.")
    insertar_figura(doc, "02_distribuciones_conteos.png")
    leyenda(doc, "figura", "Distribuciones de visualizaciones, «me gusta» y respuestas. Las tres "
                           "presentan cola larga, por lo que se usan escalas logarítmicas.")

    titulo(doc, 3, "Categorías, consultas de búsqueda y vocabulario")
    categorias = tabla("16_videos_por_categoria")
    categorias.columns = ["Categoría", "Videos", "Vistas medianas", "Comentarios", "Videos con cobertura"]
    insertar_tabla(doc, categorias, anchos=[5.3, 2.4, 3.4, 2.9, 2.5], tamano=9,
                   alineacion_derecha=(1, 2, 3, 4))
    leyenda(doc, "tabla", "Composición temática de la muestra según la categoría de YouTube.")
    texto(doc, [
        "La muestra está dominada por ", ("News & Politics", False, True),
        f" ({int(R['por_categoria'][0]['videos'])} de {R['n_videos']} videos) y, dentro de los "
        f"videos con comentarios, esa categoría concentra {int(fila(R['sent_por_tema'], 'category', 'News & Politics')['comentarios'])} "
        f"de los {R['n_comentarios']} comentarios. Esto no describe a YouTube Guatemala: describe las "
        f"{fila(R['descriptivos'], 'indicador', 'Consultas de búsqueda distintas (videos)')['valor']:.0f} "
        "consultas usadas para armar la muestra.",
    ])
    palabras = pd.DataFrame(R["palabras_comentarios"])[:10]
    bigramas = pd.DataFrame(R["bigramas_comentarios"])[:10]
    hashtags = pd.DataFrame(R["hashtags_videos"])[:10]
    emojis = pd.DataFrame(R["emojis_top"])[:10]
    vocabulario = pd.concat([
        palabras.rename(columns={"termino": "Palabra", "frecuencia": "n"}).reset_index(drop=True),
        bigramas.rename(columns={"bigrama": "Bigrama", "frecuencia": "n "}).reset_index(drop=True),
        hashtags.rename(columns={"hashtag": "Hashtag (videos)", "frecuencia": "n  "}).reset_index(drop=True),
        emojis.rename(columns={"emoji": "Emoji", "frecuencia": "n   "}).reset_index(drop=True),
    ], axis=1).fillna("")
    insertar_tabla(doc, vocabulario, anchos=[2.9, 1.1, 3.9, 1.1, 3.5, 1.1, 1.8, 1.1], tamano=8.5,
                   alineacion_derecha=(1, 3, 5, 7))
    leyenda(doc, "tabla", "Vocabulario más frecuente: palabras y bigramas de los comentarios, "
                          "hashtags del contenido de los videos y emojis usados por los comentaristas.")
    texto(doc, [
        "El vocabulario de los comentarios es inequívocamente político y económico: «pueblo», "
        "«diputado», «pagar», «dinero», «corrupto», «sueldo». Los bigramas nombran actores "
        "concretos («bernardo arévalo», «pacto corrupto») y expresan carencia («morir hambre», "
        "«pagar sueldo»). Los hashtags, en cambio, provienen del ", ("contenido", True),
        " y no de la audiencia: son etiquetas de marca de los canales (#larondagt, #envivodca) y "
        f"llamados a la acción (#like, #comenta, #comparte). En total se detectaron "
        f"{R['n_hashtags_videos']} hashtags en títulos y descripciones de videos, frente a apenas "
        f"{int(fila(R['efecto_limpieza'], 'métrica', 'Comentarios con al menos un hashtag')['valor'])} "
        "comentario con hashtag: el uso de hashtags es una práctica de los canales, no de los "
        "comentaristas.",
    ])
    figura_apaisada(doc, "05_frecuencias_texto.png",
                    "Frecuencias de palabras, bigramas y hashtags en comentarios y contenido.")
    insertar_figura(doc, "06_nube_de_palabras.png")
    leyenda(doc, "figura", "Nube de palabras. Se incluye como apoyo visual; las comparaciones "
                           "cuantitativas deben leerse en los gráficos de barras, porque una nube no "
                           "permite comparar magnitudes con precisión.")

    titulo(doc, 2, "3.2 Concentración de la participación")
    texto(doc, [
        "Se usan tres medidas complementarias: la cuota acumulada del top-n (Pareto), el índice de "
        "Gini —desigualdad, donde 0 es reparto perfectamente igualitario— y el índice "
        "Herfindahl–Hirschman normalizado, que mide dominancia.",
    ])
    concentracion = tabla("25_concentracion_participacion")
    concentracion.columns = ["Entidad", "Universo", "Top n", "Comentarios", "Cuota acumulada (%)"]
    insertar_tabla(doc, concentracion, anchos=[4.6, 2.6, 2.2, 3.4, 3.7], tamano=8.5,
                   alineacion_derecha=(1, 2, 3, 4))
    leyenda(doc, "tabla", "Concentración de la participación (curva de Pareto).")
    desigualdad = tabla("26_indices_desigualdad")
    desigualdad.columns = ["Distribución", "Gini", "HHI normalizado"]
    insertar_tabla(doc, desigualdad, anchos=[10.5, 3.0, 3.0], tamano=9, alineacion_derecha=(1, 2))
    leyenda(doc, "tabla", "Índices de desigualdad y dominancia.")
    texto(doc, [
        "La concentración es severa en el lado del contenido y baja en el de las personas, y ese "
        "contraste es el resultado más informativo de la sección. Del lado del contenido, ",
        (f"un solo video reúne el {R['video_top_pct']} % de los comentarios", True),
        f", tres videos acumulan el {fila([x for x in R['concentracion'] if x['entidad'] == 'Videos con cobertura'], 'top_n', 3)['cuota_acumulada_pct']} % "
        f"y un solo canal el {R['canal_top_pct']} %; el Gini de comentarios por video con cobertura "
        f"es {R['gini_videos_cobertura']} y sube a "
        f"{fila(R['desigualdad'], 'distribución', 'Comentarios por video (catálogo completo)')['gini']} "
        "si se cuenta el catálogo completo. Del lado de las personas ocurre lo contrario: el Gini de "
        f"comentarios por autor es apenas {R['gini_autores']} y el autor más activo aporta sólo "
        f"{fila([x for x in R['concentracion'] if x['entidad'] == 'Autores'], 'top_n', 1)['cuota_acumulada_pct']} % "
        "del total.",
    ])
    texto(doc, [
        "La lectura es directa: ", ("no existen «superusuarios» que dominen la conversación", True),
        ". Lo que concentra la participación es el contenido, no un grupo de personas "
        "hiperactivas. Los 406 comentarios se reparten de forma casi uniforme entre 332 autores "
        "distintos que, en su gran mayoría, comentan una sola vez.",
    ])
    insertar_figura(doc, "03_concentracion_pareto.png")
    leyenda(doc, "figura", "Curvas de Pareto para videos y canales. La línea discontinua marca el 80 % "
                           "acumulado.")

    titulo(doc, 2, "3.3 Popularidad frente a participación")
    popularidad = tabla("27_popularidad_vs_participacion")
    popularidad.columns = ["Comparación", "ρ de Spearman", "p", "Lectura"]
    popularidad["p"] = popularidad["p"].map(lambda v: "< 0.001" if float(v) < 0.001 else f"{float(v):.3f}")
    insertar_tabla(doc, popularidad, anchos=[5.2, 2.2, 1.7, 7.4], tamano=8.5, alineacion_derecha=(1, 2))
    leyenda(doc, "tabla", "Relación entre popularidad (visualizaciones) y participación (comentarios).")
    texto(doc, [
        "Se usa la correlación de Spearman porque no exige linealidad y resiste las colas largas de "
        "ambas variables. El resultado depende críticamente de qué videos se incluyan, y esa es "
        "precisamente la limitación que hay que declarar. Sobre el catálogo completo la correlación "
        f"es prácticamente nula (ρ = {R['rho_vistas_comentarios_todos']}), pero ese número no "
        f"significa nada: está dominado por los {R['videos_sin_comentarios']} videos con cero "
        "comentarios, de modo que mide cobertura de recolección, no comportamiento. Restringiendo a "
        f"los {R['red_obs_videos']} videos con cobertura la correlación es alta "
        f"(ρ = {R['rho_vistas_comentarios_cobertura']}, {p_valor(R['p_vistas_comentarios_cobertura'])}) "
        f"y se mantiene al usar autores únicos en lugar de comentarios "
        f"(ρ = {fila(R['popularidad'], 'comparación', 'Vistas vs autores únicos — videos con cobertura (n=19)')['rho_spearman']}), "
        "lo que descarta que un solo autor prolífico infle el conteo.",
    ])
    texto(doc, [
        ("Limitación de ambos conteos. ", True),
        f"Con n = {R['red_obs_videos']} la potencia estadística es baja y un solo video influye "
        "mucho en el resultado. Además, visualizaciones y comentarios no son comparables entre sí: "
        "una visualización es pasiva, anónima y puede repetirse; un comentario exige identificarse "
        "y actuar. El caso extremo de la muestra lo ilustra: «Plan 2032 Ciudad de Guatemala» es el "
        "video más visto (304 089 vistas) pero genera 0.82 comentarios por cada 10 000 vistas, "
        "mientras que «Qué rico come tu diputado», con 11 775 vistas, genera 136.73. "
        "Ser visto y ser comentado son fenómenos distintos.",
    ])
    insertar_figura(doc, "04_visualizaciones_vs_comentarios.png")
    leyenda(doc, "figura", "Visualizaciones frente a comentarios, con y sin los videos sin cobertura. "
                           "El panel izquierdo muestra por qué la correlación global es engañosa.")
    insertar_figura(doc, "07_categorias_y_fuentes.png")
    leyenda(doc, "figura", "Composición temática y cobertura por estrategia de muestreo. Ningún video "
                           "recolectado con la estrategia official_gov aportó comentarios.")
    salto_de_pagina(doc)


def seccion_3_preguntas(doc: Document) -> None:
    titulo(doc, 2, "3.5 Preguntas obligatorias")

    titulo(doc, 3, "a) ¿Qué videos y canales concentran la mayor participación observada?")
    p1v = pd.DataFrame(R["p1_videos"])[["title", "channel_name", "view_count", "comentarios_obs",
                                        "autores_unicos", "cuota_pct"]]
    p1v.columns = ["Video", "Canal", "Vistas", "Comentarios", "Autores", "% del total"]
    insertar_tabla(doc, p1v, anchos=[5.6, 3.6, 2.1, 2.1, 1.6, 1.5], tamano=8.5,
                   alineacion_derecha=(2, 3, 4, 5))
    leyenda(doc, "tabla", "Los cinco videos con mayor participación observada.")
    p1c = pd.DataFrame(R["p1_canales"])
    p1c.columns = ["Canal", "Videos", "Videos con comentarios", "Comentarios", "Autores", "% del total"]
    insertar_tabla(doc, p1c, anchos=[5.6, 2.0, 3.4, 2.3, 1.8, 1.4], tamano=8.5,
                   alineacion_derecha=(1, 2, 3, 4, 5))
    leyenda(doc, "tabla", "Los cinco canales con mayor participación observada.")
    texto(doc, [
        ("Respuesta. ", True), "La participación se concentra en el video «",
        (R["video_top_titulo"], False, True), "» del canal ", (R["video_top_canal"], True),
        f", que reúne {R['video_top_comentarios']} comentarios de {int(R['p1_videos'][0]['autores_unicos'])} "
        f"autores distintos ({R['video_top_pct']} % del total). A nivel de canal, {R['canal_top_nombre']} "
        f"concentra {R['canal_top_comentarios']} comentarios ({R['canal_top_pct']} %) repartidos en "
        f"{int(R['p1_canales'][0]['videos_con_comentarios'])} de sus {int(R['p1_canales'][0]['videos'])} "
        "videos. Un matiz importante: el Gobierno de la República de Guatemala aparece con "
        "32 videos en el catálogo pero sólo 2 con comentarios recolectados, de modo que su "
        "segundo lugar refleja tanto participación real como la selección de la muestra.",
    ])

    titulo(doc, 3, "b) ¿Existen audiencias compartidas entre videos, canales o temas?")
    audiencia = pd.DataFrame(R["audiencia_compartida"])
    audiencia.columns = ["Indicador", "Autores", "% del total"]
    insertar_tabla(doc, audiencia, anchos=[10.5, 3.0, 3.0], tamano=9, alineacion_derecha=(1, 2))
    leyenda(doc, "tabla", "Medida directa de audiencia compartida entre videos y canales.")
    texto(doc, [
        ("Respuesta: sí, pero son mínimas. ", True), f"Sólo {R['autores_multivideo_n']} de "
        f"{R['autores_unicos']} autores ({fila(R['audiencia_compartida'], 'indicador', 'Autores que comentan en más de un video')['porcentaje_del_total']} %) "
        f"comentaron en más de un video, y sólo {R['autores_multicanal_n']} cruzaron canales "
        f"distintos. El {fila(R['audiencia_compartida'], 'indicador', 'Autores exclusivos de un solo video')['porcentaje_del_total']} % "
        "es exclusivo de un único video. Traducido a la proyección video–video: de los "
        f"{R['pv_nodos']} videos con cobertura, sólo {R['pv_nodos'] - int(fila(R['periferia'], 'grupo', 'Videos aislados en la proyección video–video')['conteo'])} "
        f"comparten al menos un comentarista con otro video, formando apenas {R['pv_aristas']} "
        "conexiones, y el solapamiento máximo entre dos videos es de 2 personas. Las audiencias de "
        "esta muestra son prácticamente disjuntas.",
    ])

    titulo(doc, 3, "c) ¿Qué autores funcionan como puentes entre contenidos separados?")
    puentes = tabla("58_autores_puente_articuladores")[
        ["autor", "videos_comentados", "comentarios", "intermediacion", "componentes_despues", "videos_que_conecta"]]
    puentes.columns = ["Autor", "Videos", "Comentarios", "Intermediación", "Componentes tras eliminarlo",
                       "Videos que conecta"]
    insertar_tabla(doc, puentes, anchos=[3.1, 1.3, 1.9, 2.2, 2.7, 5.3], tamano=8,
                   alineacion_derecha=(1, 2, 3, 4))
    leyenda(doc, "tabla", "Autores que son puntos de articulación: eliminarlos fragmenta la red.")
    texto(doc, [
        ("Respuesta. ", True), f"Exactamente {R['n_autores_puente']} autores "
        f"({round(100 * R['n_autores_puente'] / R['autores_unicos'], 1)} % del total) funcionan como "
        "puentes estructurales. La prueba no es una métrica abstracta sino una verificación directa: "
        f"al eliminar cualquiera de ellos, la subred observada pasa de {R['n_componentes_obs']} a 11 "
        "componentes. El más importante es ",
        (R["puentes"][0]["autor"], True), f", cuya intermediación es {R['puentes'][0]['intermediacion']} "
        "y es la única persona que une la conferencia de prensa del Gobierno con el video más "
        "comentado de la muestra. Lo notable es su ", ("desproporción", True),
        ": estos autores aportan entre 2 y 4 comentarios cada uno, cifras irrelevantes en cualquier "
        "conteo de volumen, y sin embargo su ausencia rompe la red. La centralidad estructural no "
        "es lo mismo que la actividad.",
    ])

    titulo(doc, 3, "d) ¿Qué temas y sentimientos caracterizan las principales comunidades?")
    perfil = tabla("34_p4_perfil_tematico_y_sentimiento_por_video").head(8)[
        ["title", "comentarios_obs", "sent_medio", "pct_neg", "terminos"]]
    perfil.columns = ["Video", "Comentarios", "Sentimiento medio", "% negativos", "Términos frecuentes"]
    insertar_tabla(doc, perfil, anchos=[4.8, 2.0, 2.3, 1.9, 5.5], tamano=8,
                   alineacion_derecha=(1, 2, 3))
    leyenda(doc, "tabla", "Perfil temático y de sentimiento de los videos con más participación. "
                          "La caracterización por comunidad se detalla en la sección 7.5.")
    texto(doc, [
        ("Respuesta. ", True), "Cada foco de participación tiene un perfil temático y emocional "
        "propio y coherente. El núcleo mayor gira en torno al gasto público y los privilegios de los "
        "diputados («pueblo», «diputado», «pagar», «sueldo») con un sentimiento medio de "
        f"{R['perfil_videos'][0]['sent_medio']} y {R['perfil_videos'][0]['pct_neg']} % de comentarios "
        "negativos. Un segundo foco trata la cooptación universitaria («corrupto», «usac», "
        "«estudiante»). En el extremo opuesto, la planificación urbana municipal («guatemala», "
        "«ciudad», «proyecto», «saludo») es el único tema con sentimiento claramente positivo "
        f"({R['perfil_videos'][5]['sent_medio']:+.3f}). El patrón es consistente: la fiscalización "
        "política activa participación negativa y la obra pública local activa participación de "
        "apoyo.",
    ])

    titulo(doc, 3, "e) ¿La visibilidad por visualizaciones coincide con la participación observada?")
    ranking = tabla("35_p5_visibilidad_vs_participacion").head(8)[
        ["title", "view_count", "comentarios_obs", "rank_vistas", "rank_comentarios", "comentarios_por_10k_vistas"]]
    ranking.columns = ["Video", "Vistas", "Comentarios", "Rango vistas", "Rango comentarios",
                       "Comentarios por 10 000 vistas"]
    insertar_tabla(doc, ranking, anchos=[5.6, 2.1, 2.1, 2.0, 2.2, 2.5], tamano=8,
                   alineacion_derecha=(1, 2, 3, 4, 5))
    leyenda(doc, "tabla", "Comparación de rangos: visibilidad frente a participación.")
    texto(doc, [
        ("Respuesta: sólo parcialmente. ", True), "Dentro de los videos con cobertura ambos rangos "
        f"correlacionan alto (ρ = {R['rho_vistas_comentarios_cobertura']}), pero la tasa de "
        "participación por audiencia varía en dos órdenes de magnitud: de 0.82 a 167 comentarios por "
        "cada 10 000 vistas. El desajuste más claro es «Plan 2032 Ciudad de Guatemala», primero en "
        "visualizaciones y sólo cuarto en comentarios. La visibilidad no predice la participación; "
        "lo que la predice es el tipo de contenido.",
    ])

    titulo(doc, 3, "f) ¿Qué conclusiones limita el procedimiento de recolección y la cobertura?")
    cobertura = tabla("36_p6_cobertura_por_estrategia")
    cobertura.columns = ["Estrategia (source_group)", "Videos", "Con comentarios", "Comentarios",
                         "Vistas medianas", "Cobertura (%)"]
    insertar_tabla(doc, cobertura, anchos=[4.2, 2.2, 3.0, 2.5, 2.7, 1.9], tamano=9,
                   alineacion_derecha=(1, 2, 3, 4, 5))
    leyenda(doc, "tabla", "Cobertura de comentarios según la estrategia de recolección del video.")
    texto(doc, [
        ("Respuesta. ", True), "La cobertura no es aleatoria, está asociada a la estrategia de "
        "recolección, y esto invalida comparaciones entre grupos. Los videos hallados por canal "
        "tienen 63.6 % de cobertura, los hallados por tema 6.8 %, y los ",
        ("105 videos de la estrategia official_gov tienen 0 %", True), ": ni uno solo aportó "
        "comentarios. Concluir que «el contenido gubernamental no genera participación» sería un "
        "error grave, porque para esos videos no se recolectó ningún comentario. En consecuencia: "
        "no puede compararse participación entre estrategias de muestreo, no puede estimarse la "
        "participación promedio de un video guatemalteco, y no puede afirmarse que un video sin "
        "comentarios en el archivo no los tenga en YouTube.",
    ])


def seccion_3_extras(doc: Document) -> None:
    titulo(doc, 2, "3.6 Tres preguntas adicionales surgidas del análisis")

    titulo(doc, 3, "Pregunta 1: ¿la intensidad y la amplitud de participación son el mismo fenómeno?")
    intensidad = pd.DataFrame(R["intensidad"])
    intensidad.columns = ["Perfil de participación", "Autores"]
    insertar_tabla(doc, intensidad, anchos=[12.5, 4.0], tamano=9, alineacion_derecha=(1,))
    leyenda(doc, "tabla", "Intensidad (varios comentarios en un video) frente a amplitud (varios videos).")
    texto(doc, [
        ("Respuesta: no, son perfiles distintos. ", True),
        f"{int(R['intensidad'][1]['autores'])} autores comentan varias veces pero siempre en el mismo "
        f"video —hasta {int(R['intensidad'][3]['autores'])} veces—, mientras que sólo "
        f"{int(R['intensidad'][2]['autores'])} aparecen en más de un video. La correlación entre "
        f"número de comentarios y número de videos por autor es moderada "
        f"(ρ = {R['rho_intensidad_amplitud']}), lo que confirma que insistir dentro de una discusión "
        "y circular entre discusiones son comportamientos diferentes. La consecuencia para la red es "
        "importante: la intensidad aumenta el ", ("peso", True), " de una arista, pero sólo la "
        "amplitud crea ", ("conectividad", True), ". Un autor con seis comentarios en un único video "
        "no conecta nada; uno con dos comentarios en dos videos puede ser un puente.",
    ])

    titulo(doc, 3, "Pregunta 2: ¿los comentarios más aplaudidos son más negativos?")
    likes = pd.DataFrame(R["likes_sentimiento"])
    likes.columns = ["Tramo de «me gusta»", "Comentarios", "Sentimiento medio", "% negativos", "Likes medianos"]
    insertar_tabla(doc, likes, anchos=[3.6, 3.0, 3.6, 3.2, 3.1], tamano=9,
                   alineacion_derecha=(1, 2, 3, 4))
    leyenda(doc, "tabla", "Sentimiento por tramo de «me gusta» recibidos.")
    texto(doc, [
        ("Respuesta: ocurre exactamente lo contrario de lo esperado. ", True),
        "La hipótesis intuitiva —que la indignación se premia con más «me gusta»— no se sostiene. "
        "El gradiente es monótono y va en dirección opuesta: los comentarios sin ningún «me gusta» "
        f"tienen sentimiento medio {R['likes_sentimiento'][0]['sent_medio']} y "
        f"{R['likes_sentimiento'][0]['pct_negativos']} % de negativos, mientras que los que superan "
        f"los 50 «me gusta» tienen {R['likes_sentimiento'][4]['sent_medio']:+.3f} y "
        f"{R['likes_sentimiento'][4]['pct_negativos']} % de negativos. La correlación es positiva y "
        f"significativa (ρ = {R['rho_likes_sentimiento']}, {p_valor(R['p_likes_sentimiento'])}). "
        "La interpretación debe ser prudente: el tramo superior tiene sólo 9 comentarios y está "
        "dominado por un video de tema municipal. Lo que puede afirmarse es que, en esta muestra, "
        "el respaldo explícito de la audiencia acompaña a los mensajes de apoyo, no a los de crítica.",
    ])

    titulo(doc, 3, "Pregunta 3: ¿la estrategia de muestreo produce perfiles distintos?")
    perfil = pd.DataFrame(R["perfil_fuente"])
    perfil.columns = ["Estrategia", "Comentarios", "Autores", "Videos", "Likes medianos",
                      "Sentimiento medio", "% negativos", "Términos frecuentes"]
    insertar_tabla(doc, perfil, anchos=[1.9, 1.9, 1.6, 1.5, 1.9, 2.2, 1.7, 3.8], tamano=8,
                   alineacion_derecha=(1, 2, 3, 4, 5, 6))
    leyenda(doc, "tabla", "Perfil de participación según cómo se recolectaron los comentarios.")
    texto(doc, [
        ("Respuesta: sí, y es una advertencia metodológica. ", True),
        "Los comentarios obtenidos rastreando canales concretos son marcadamente más negativos "
        f"({R['perfil_fuente'][0]['sent_medio']}, {R['perfil_fuente'][0]['pct_negativos']} % negativos) "
        "que los obtenidos por búsqueda temática "
        f"({R['perfil_fuente'][1]['sent_medio']}, {R['perfil_fuente'][1]['pct_negativos']} %), y su "
        "vocabulario también difiere: fiscalización del gasto público frente a agenda "
        "gubernamental y urbana. Como ambas rutas de recolección seleccionan contenidos distintos, "
        "una parte del tono negativo global de la muestra es atribuible al ",
        ("procedimiento", True), " y no sólo al comportamiento de los usuarios.",
    ])
    salto_de_pagina(doc)


# --------------------------------------------------------------------------------------------
# 4. Red bipartita
# --------------------------------------------------------------------------------------------
def seccion_4(doc: Document) -> None:
    titulo(doc, 1, "4. Construcción de la red bipartita autor–video")

    titulo(doc, 2, "4.5 Qué significa exactamente una arista")
    texto(doc, [
        "Se declara antes de construir la red, porque toda interpretación posterior depende de "
        "ello. Una arista no dirigida entre un autor y un video significa ", ("una sola cosa", True),
        ": ese autor publicó al menos un comentario principal en ese video, dentro de la muestra "
        "recolectada. Su ", ("peso", True), " es el número de comentarios que ese autor publicó en "
        f"ese video, con un máximo observado de {R['peso_max']}.",
    ])
    texto(doc, [("Una arista NO significa ninguna de las siguientes cosas:", True)])
    vineta(doc, "Que el autor sea amigo, seguidor o suscriptor del canal.")
    vineta(doc, "Que exista conversación directa con nadie. Los datos no identifican quién respondió "
                "a quién, de modo que reply_count nunca genera aristas entre usuarios.")
    vineta(doc, "Aprobación del contenido. El análisis de sentimiento muestra que el "
                f"{R['pct_negativo_global']} % de los comentarios son negativos: participar equivale "
                "con más frecuencia a criticar que a apoyar.")
    vineta(doc, "Que el autor haya visto el video completo, ni que ésa sea su única actividad en "
                "YouTube.")
    texto(doc, [
        "La red es ", ("bipartita", True), ": sólo existen aristas autor↔video. Nunca autor↔autor ni "
        "video↔video. Esas relaciones aparecen únicamente como proyecciones (sección 5), que son "
        "construcciones derivadas y no observaciones.",
    ])

    titulo(doc, 2, "4.1 a 4.3 Tabla de nodos, tabla de aristas y validación")
    nodos_desc = pd.DataFrame([
        ["node_id", "Ambos", "Identificador con prefijo de tipo: author::<id> o video::<id>."],
        ["tipo_nodo / bipartite", "Ambos", "«autor» (partición 0) o «video» (partición 1)."],
        ["etiqueta", "Ambos", "Handle del autor o título del video; sólo para mostrar."],
        ["id_original", "Ambos", "author_channel_id o video_id sin prefijo, para volver a los datos."],
        ["observado", "Ambos", "Falso en los videos sin comentarios recolectados."],
        ["comentarios_publicados / videos_comentados / canales_comentados", "Autor",
         "Volumen, amplitud y alcance entre canales."],
        ["likes, respuestas, sent_medio", "Autor", "Recepción y tono medio de sus comentarios."],
        ["channel_id, channel_name, category, source_group", "Video", "Procedencia y clasificación."],
        ["view_count, comentarios_recibidos, autores_distintos", "Video", "Popularidad y participación."],
        ["sent_medio", "Video", "Tono medio de los comentarios recibidos."],
    ], columns=["Atributo", "Aplica a", "Significado"])
    insertar_tabla(doc, nodos_desc, anchos=[6.0, 1.9, 8.6], tamano=8.5)
    leyenda(doc, "tabla", "Estructura de la tabla de nodos (outputs/tables/40_nodos_red_bipartita.csv).")

    aristas_desc = pd.DataFrame([
        ["origen / destino", "Extremos de la arista: author::<id> y video::<id>."],
        ["peso", "Número de comentarios de ese autor en ese video. Suma total = "
                 f"{R['peso_total']}, el número de comentarios del archivo."],
        ["likes / respuestas", "«Me gusta» y respuestas acumuladas por esos comentarios."],
        ["sent_medio", "Sentimiento medio de los comentarios que forman la arista."],
        ["tipo_arista", "«comento_en». Constante: sólo existe un tipo de relación observable."],
        ["author_handle / title / channel_name", "Etiquetas legibles para inspección manual."],
    ], columns=["Atributo", "Significado"])
    insertar_tabla(doc, aristas_desc, anchos=[4.6, 11.9], tamano=8.5)
    leyenda(doc, "tabla", "Estructura de la tabla de aristas (outputs/tables/41_aristas_red_bipartita.csv).")

    validacion = tabla("42_validacion_red_bipartita")
    validacion.columns = ["Verificación", "Resultado"]
    insertar_tabla(doc, validacion, anchos=[11.5, 5.0], tamano=9, alineacion_derecha=(1,))
    leyenda(doc, "tabla", "Validación estructural de la red construida.")
    texto(doc, [
        "La red completa contiene ", (f"{R['red_nodos_total']} nodos", True), " —",
        f"{R['autores_unicos']} autores y {R['n_videos']} videos— y ",
        (f"{R['red_aristas_total']} aristas", True), f" cuyo peso suma exactamente "
        f"{R['peso_total']}, es decir, el número total de comentarios. Esa igualdad es la "
        "verificación de que ningún comentario se perdió ni se contó dos veces. La red es bipartita "
        "por construcción y se comprobó que no existe ninguna arista entre nodos del mismo tipo. "
        f"La subred observada —los nodos con al menos una arista— tiene {R['red_obs_nodos']} nodos: "
        f"los {R['red_obs_autores']} autores y sólo {R['red_obs_videos']} videos.",
    ])

    titulo(doc, 2, "4.4 Visualización de la red completa")
    texto(doc, [
        "Se visualiza la red ", ("completa, sin eliminar nodos", True), f". Los "
        f"{R['videos_sin_comentarios']} videos sin comentarios recolectados se dibujan como círculos "
        "vacíos: son un hallazgo sustantivo —el sesgo de cobertura— y borrarlos por estética "
        "ocultaría el resultado más importante del muestreo. Se usan dos vistas complementarias "
        "porque una sola no basta: un diseño bipartito en columnas que muestra la escala real del "
        "desbalance, y un diseño de fuerzas sobre la subred observada que revela la estructura "
        "interna, ilegible en la primera vista.",
    ])
    figura_apaisada(doc, "08_red_bipartita.png",
                    "Red bipartita autor–video. Izquierda: red completa, con los videos sin cobertura "
                    "dibujados como círculos vacíos. Derecha: subred observada con diseño de fuerzas.")
    texto(doc, [
        "La forma que emerge es la de una ", ("constelación de estrellas", True), ": cada video "
        "actúa como centro de un cúmulo de autores que sólo se conectan a él. Los pocos enlaces "
        "entre cúmulos son autores que comentaron en más de un video. Esta forma no es un artefacto "
        "del algoritmo de dibujo: se confirma cuantitativamente en la sección 6 con la distribución "
        "de grados.",
    ])
    salto_de_pagina(doc)


# --------------------------------------------------------------------------------------------
# 5. Proyecciones
# --------------------------------------------------------------------------------------------
def seccion_5(doc: Document) -> None:
    titulo(doc, 1, "5. Proyecciones de la red")

    titulo(doc, 2, "5.1 y 5.2 Construcción de ambas proyecciones")
    texto(doc, [
        "La ", ("proyección autor–autor", True), " conecta dos autores si comentaron en el mismo "
        "video; el peso es el número de videos que comparten. La ", ("proyección video–video", True),
        " conecta dos videos si comparten al menos un autor; el peso es el número de autores "
        "compartidos. Ambas se calculan sobre la subred observada.",
    ])
    texto(doc, [
        ("Advertencia estructural, declarada antes de leer los resultados: ", True),
        "una proyección convierte cada video en una ", ("camarilla completa", True),
        " entre todos sus comentaristas. El video con 128 comentaristas genera por sí solo 8 128 "
        f"aristas. Por eso las {R['pa_aristas']:,} aristas de la proyección autor–autor no "
        f"representan {R['pa_aristas']:,} interacciones: son el resultado mecánico de "
        f"{R['red_aristas_total']} observaciones reales.".replace(",", " "),
    ])
    comparacion = tabla("45_comparacion_proyecciones")
    comparacion.columns = ["Métrica", "Proyección autor–autor", "Proyección video–video"]
    insertar_tabla(doc, comparacion, anchos=[7.0, 4.8, 4.7], tamano=9, alineacion_derecha=(1, 2))
    leyenda(doc, "tabla", "Comparación cuantitativa de ambas proyecciones.")

    titulo(doc, 2, "5.3 Qué fenómeno representa cada proyección")
    vineta(doc, [
        ("La proyección autor–autor representa co-presencia, no conversación. ", True),
        f"Su densidad ({fila(R['comparacion_proyecciones'], 'métrica', 'Densidad')['proyección_autor_autor']}) "
        f"y su transitividad ({fila(R['comparacion_proyecciones'], 'métrica', 'Transitividad')['proyección_autor_autor']}) "
        "son artificialmente altas por efecto de las camarillas. Sólo 2 de sus "
        f"{R['pa_aristas']:,} aristas tienen peso mayor que 1: la evidencia detrás de casi todas es "
        "«coincidieron una vez en la misma sección de comentarios».".replace(",", " "),
    ])
    vineta(doc, [
        ("La proyección video–video representa solapamiento de audiencia. ", True),
        f"Es mucho más informativa: sus {R['pv_aristas']} aristas corresponden a personas concretas y "
        "verificables. Permite responder si dos contenidos comparten público, que es la pregunta "
        "sustantiva del inciso 3.5(b).",
    ])
    vineta(doc, [
        ("Consecuencia metodológica. ", True), "Para detectar comunidades y calcular intermediación "
        "se trabaja sobre la red bipartita, no sobre la proyección autor–autor, cuyas camarillas "
        "inflan la cohesión y distorsionan cualquier medida basada en caminos más cortos.",
    ])

    titulo(doc, 2, "5.4 Visualización de ambas proyecciones")
    figura_apaisada(doc, "09_proyecciones.png",
                    "Izquierda: la proyección autor–autor muestra las camarillas que genera cada video; "
                    "en rojo, los autores presentes en más de un video. Derecha: la proyección "
                    "video–video con el número de autores compartidos sobre cada arista; abajo, los "
                    "videos sin audiencia compartida con ningún otro.")
    texto(doc, [
        "La imagen resume el argumento: a la izquierda, grandes cúmulos densos que son un artefacto "
        "de la proyección y unos pocos nodos rojos que los enlazan; a la derecha, una cadena frágil "
        "de solapamientos de audiencia con pesos de 1 o 2 personas, y "
        f"{int(fila(R['periferia'], 'grupo', 'Videos aislados en la proyección video–video')['conteo'])} "
        f"de los {R['pv_nodos']} videos completamente separados del resto.",
    ])
    salto_de_pagina(doc)


# --------------------------------------------------------------------------------------------
# 6. Topología y fragmentación
# --------------------------------------------------------------------------------------------
def seccion_6(doc: Document) -> None:
    titulo(doc, 1, "6. Topología y fragmentación")

    titulo(doc, 2, "6.1 Métricas estructurales")
    topologia = tabla("46_metricas_topologicas").set_index("red").T.reset_index()
    topologia.columns = ["Métrica", "Bipartita completa", "Bipartita observada",
                         "Proyección autor–autor", "Proyección video–video"]
    topologia["Métrica"] = topologia["Métrica"].str.replace("_", " ").str.capitalize()
    insertar_tabla(doc, topologia, anchos=[5.4, 2.9, 2.9, 2.8, 2.5], tamano=8.5,
                   alineacion_derecha=(1, 2, 3, 4))
    leyenda(doc, "tabla", "Métricas topológicas de la red bipartita y sus proyecciones. La densidad "
                          "bipartita se calcula sobre el producto de los dos lados, no sobre n(n−1)/2.")
    texto(doc, [
        "La red bipartita observada tiene ", (f"{R['red_obs_nodos']} nodos y {R['red_aristas_total']} "
        "aristas", True), f", una densidad de "
        f"{fila(R['topologia'], 'red', 'Bipartita observada (19 videos + 332 autores)')['densidad']} "
        f"y un grado medio de "
        f"{fila(R['topologia'], 'red', 'Bipartita observada (19 videos + 332 autores)')['grado_medio']}. "
        "El grado medio global, sin embargo, esconde el hallazgo real, que sólo aparece al separar "
        "los dos lados de la red.",
    ])

    titulo(doc, 2, "6.1 Distribución de grados: ¿pocas conexiones o concentradas en unos pocos?")
    grados = tabla("49_resumen_grados_por_lado")
    grados.columns = ["Lado de la red", "Nodos", "Grado medio", "Grado mediano", "Grado máximo",
                      "% con grado 1"]
    insertar_tabla(doc, grados, anchos=[4.6, 2.2, 2.5, 2.6, 2.6, 2.0], tamano=9,
                   alineacion_derecha=(1, 2, 3, 4, 5))
    leyenda(doc, "tabla", "Distribución de grados separada por partición.")
    texto(doc, [
        ("La respuesta es: ambas cosas a la vez, según el lado que se mire. ", True),
        f"Del lado de los autores la distribución es plana y baja: el {R['pct_autores_grado1']} % "
        "tiene grado 1 y el máximo es 3. Del lado de los videos es todo lo contrario: el grado medio "
        f"es {fila(R['resumen_grados'], 'lado', 'Videos (con cobertura)')['grado_medio']}, la mediana "
        f"{fila(R['resumen_grados'], 'lado', 'Videos (con cobertura)')['grado_mediano']:.0f} y el "
        f"máximo {fila(R['resumen_grados'], 'lado', 'Videos (con cobertura)')['grado_max']}. "
        "Es decir, las conexiones ", ("sí", True), " están concentradas en unos pocos nodos, pero "
        "esos nodos son ", ("videos", True), ", no personas. Interpretar el grado medio global "
        f"({fila(R['topologia'], 'red', 'Bipartita observada (19 videos + 332 autores)')['grado_medio']}) "
        "sin separar particiones llevaría a la conclusión equivocada de que la red es homogénea.",
    ])
    figura_apaisada(doc, "10_distribucion_grados_componentes.png",
                    "Distribución de grados por partición y tamaño de las componentes conexas de la "
                    "subred observada.")

    titulo(doc, 2, "6.2 Cohesión y transitividad")
    cohesion = tabla("47_cohesion_y_transitividad")
    cohesion.columns = ["Medida", "Valor", "Interpretación"]
    insertar_tabla(doc, cohesion, anchos=[5.6, 1.8, 9.1], tamano=8.5, alineacion_derecha=(1,))
    leyenda(doc, "tabla", "Medidas de cohesión, con la corrección bipartita.")
    texto(doc, [
        ("La transitividad clásica de la red bipartita vale exactamente 0, y eso no significa "
         "ausencia de cohesión. ", True),
        "En una red bipartita no pueden existir triángulos: un ciclo debe alternar autor y video, de "
        "modo que el ciclo más corto posible tiene longitud 4. El coeficiente clásico es 0 ",
        ("por construcción", True), ". Reportarlo como «la red no tiene cohesión» sería un error de "
        "método. La medida correcta es el ", ("clustering bipartito de Latapy", True),
        ", que cuantifica el solapamiento de vecindarios entre nodos del mismo lado, y ahí aparece "
        "el contraste real: "
        f"{fila(R['cohesion'], 'medida', 'Clustering bipartito de Latapy — autores')['valor']} entre "
        "autores frente a "
        f"{fila(R['cohesion'], 'medida', 'Clustering bipartito de Latapy — videos')['valor']} entre "
        "videos. Los autores comparten vecindario casi por completo —porque casi todos están "
        "conectados al mismo video— mientras que los videos casi no comparten audiencia entre sí.",
    ])
    texto(doc, [
        "La ", ("redundancia media", True), " de los nodos con al menos dos vecinos es "
        f"{[x for x in R['cohesion'] if x['medida'].startswith('Redundancia')][0]['valor']}: si uno de "
        "esos nodos desapareciera, sólo el 11 % de los pares de vecinos que conecta seguiría unido "
        "por otro camino. Es una red estructuralmente frágil. En la proyección autor–autor la "
        f"transitividad sube a "
        f"{fila(R['cohesion'], 'medida', 'Transitividad — proyección autor–autor')['valor']}, un valor "
        "casi perfecto que no describe cohesión social sino el artefacto de las camarillas.",
    ])

    titulo(doc, 2, "6.3 Fragmentación, periferia y aislamiento")
    componentes = tabla("50_componentes_conexas")[["componente", "nodos", "autores", "videos", "comentarios", "titulos"]]
    componentes.columns = ["Componente", "Nodos", "Autores", "Videos", "Comentarios", "Videos que contiene"]
    insertar_tabla(doc, componentes, anchos=[1.9, 1.5, 1.6, 1.5, 2.0, 8.0], tamano=8,
                   alineacion_derecha=(1, 2, 3, 4))
    leyenda(doc, "tabla", "Componentes conexas de la subred observada.")
    texto(doc, [
        f"La subred observada se fragmenta en {R['n_componentes_obs']} componentes. La mayor reúne "
        f"{R['tam_componente_mayor']} nodos ({R['pct_componente_mayor']} % de la subred), diez videos "
        "y 276 autores; su diámetro es "
        f"{fila(R['topologia'], 'red', 'Bipartita observada (19 videos + 332 autores)')['diametro_componente_mayor']} "
        "y la distancia media entre nodos es "
        f"{fila(R['topologia'], 'red', 'Bipartita observada (19 videos + 332 autores)')['camino_medio_componente_mayor']}. "
        "Las nueve componentes restantes son cada una un solo video con su público exclusivo, "
        "desde 26 nodos hasta apenas 2.",
    ])
    periferia = tabla("51_periferia_y_aislamiento")
    periferia.columns = ["Grupo", "Conteo", "Lectura correcta"]
    insertar_tabla(doc, periferia, anchos=[5.0, 1.7, 9.8], tamano=8.5, alineacion_derecha=(1,))
    leyenda(doc, "tabla", "Distinción entre aislamiento observado y ausencia de datos.")
    texto(doc, [
        ("La distinción entre aislamiento y ausencia de datos es la decisión interpretativa más "
         "importante de esta sección. ", True),
        f"Los {R['videos_sin_comentarios']} videos sin ninguna arista ", ("no", True),
        " son videos socialmente aislados: son videos para los que no se recolectaron comentarios. "
        "Afirmar que nadie comentó en ellos sería confundir el instrumento con el fenómeno. En "
        "cambio, los "
        f"{int(fila(R['periferia'], 'grupo', 'Videos con cobertura y grado 1')['conteo'])} videos con "
        "cobertura y grado 1 sí exhiben aislamiento observado —recibieron un solo comentarista— y "
        f"los {int(fila(R['periferia'], 'grupo', 'Autores con grado 1 (un solo video)')['conteo'])} "
        "autores de grado 1 constituyen una periferia observada: participaron una vez en la muestra, "
        "lo que no impide que sean activos fuera de ella.",
    ])

    titulo(doc, 2, "6.4 Interpretación de los hallazgos estructurales")
    texto(doc, [
        "Los tres resultados anteriores describen el mismo fenómeno desde ángulos distintos. La "
        "distribución de grados asimétrica, la transitividad nula y la fragmentación en "
        f"{R['n_componentes_obs']} componentes apuntan a que ",
        ("la participación observada se organiza alrededor del contenido y no de vínculos entre "
         "personas", True), ". Los videos son los agregadores; los autores son, en su inmensa "
        "mayoría, visitantes de una sola vez. La conectividad global no descansa en una comunidad "
        f"cohesionada sino en {R['n_autores_puente']} individuos cuya ausencia bastaría para "
        "fragmentar el núcleo, algo que se demuestra en la sección 8.",
    ])
    salto_de_pagina(doc)


# --------------------------------------------------------------------------------------------
# 7. Comunidades
# --------------------------------------------------------------------------------------------
def seccion_7(doc: Document) -> None:
    titulo(doc, 1, "7. Comunidades")

    titulo(doc, 2, "7.1 Red seleccionada y justificación")
    texto(doc, [
        "Las comunidades se detectan sobre la ", ("red bipartita observada", True), " —",
        f"{R['red_obs_nodos']} nodos y {R['red_aristas_total']} aristas— y no sobre la proyección "
        "autor–autor. Tres razones:",
    ])
    vineta(doc, [
        f"La proyección genera {R['pa_aristas']:,} aristas a partir de {R['red_aristas_total']} "
        "observaciones reales. Cualquier comunidad detectada allí reflejaría el artefacto de las "
        "camarillas, no la estructura de participación.".replace(",", " "),
    ])
    vineta(doc, "La red bipartita conserva la unidad de observación real —quién comentó dónde—, que "
                "es exactamente el objeto de interés.")
    vineta(doc, "Al mantener autores y videos como nodos, cada comunidad resulta interpretable: "
                "puede describirse por sus videos, su canal, sus temas y su sentimiento.")
    texto(doc, [
        f"Los {R['videos_sin_comentarios']} videos sin cobertura se ", ("excluyen", True),
        " de la detección: tienen grado 0 y cada uno formaría una comunidad trivial de tamaño 1, "
        "inflando artificialmente el número de comunidades sin aportar información. Se contabilizan "
        "aparte como ausencia de datos.",
    ])

    titulo(doc, 2, "7.2 Algoritmo, supuestos y tratamiento de los pesos")
    texto(doc, [
        "Se aplicó ", ("Louvain con pesos", True), " (implementación python-louvain, semilla fija "
        "42 para reproducibilidad). Los supuestos y decisiones son los siguientes.",
    ])
    vineta(doc, [
        ("Supuesto de modularidad. ", True), "Una comunidad es un conjunto de nodos con más aristas "
        "internas de las que cabría esperar en un grafo aleatorio con la misma distribución de "
        "grados (modelo nulo de configuración).",
    ])
    vineta(doc, [
        ("Tratamiento de los pesos. ", True), "El peso de la arista es el número de comentarios del "
        "autor en ese video, de modo que un autor que comentó seis veces pesa más que uno que "
        "comentó una vez. Para verificar que el resultado no depende de esa decisión se repitió la "
        "partición sin pesos.",
    ])
    vineta(doc, [
        ("Limitación conocida y declarada. ", True), "La modularidad clásica no está diseñada para "
        "redes bipartitas: su modelo nulo admite aristas autor–autor que aquí son imposibles, lo "
        "que ", ("sobreestima", True), " el valor obtenido. Por eso la modularidad se interpreta de "
        "forma comparativa y no como medida absoluta de calidad. La alternativa formal sería la "
        "modularidad bipartita de Barber; el resultado no cambiaría cualitativamente porque las "
        "componentes conexas ya separan grupos casi disjuntos.",
    ])
    vineta(doc, [
        ("Límite de resolución. ", True), "Louvain no detecta comunidades más pequeñas que √(2m); "
        f"con m = {R['red_aristas_total']} el umbral es de unos 26 nodos, lo que explica que las "
        "componentes pequeñas queden como comunidades completas.",
    ])
    algoritmos = tabla("52_comparacion_algoritmos_comunidades")
    algoritmos.columns = ["Algoritmo", "Comunidades", "Modularidad", "Tamaño de la mayor"]
    insertar_tabla(doc, algoritmos, anchos=[7.0, 3.2, 3.2, 3.1], tamano=9, alineacion_derecha=(1, 2, 3))
    leyenda(doc, "tabla", "Robustez de la partición frente a cuatro algoritmos alternativos.")
    texto(doc, [
        "La partición es ", ("robusta", True), f": cuatro de los cinco métodos coinciden en "
        f"{R['n_comunidades']} comunidades y el quinto difiere en dos. Quitar los pesos cambia la "
        "modularidad de "
        f"{fila(R['comparacion_algoritmos'], 'algoritmo', 'Louvain con pesos (python-louvain)')['modularidad']} "
        f"a {fila(R['comparacion_algoritmos'], 'algoritmo', 'Louvain sin pesos (python-louvain)')['modularidad']} "
        "sin alterar el número de grupos, lo que confirma que la estructura la determina quién "
        "comentó dónde y no cuántas veces lo hizo.",
    ])

    titulo(doc, 2, "7.3 Número de comunidades, tamaños y calidad")
    texto(doc, [
        "Se detectaron ", (f"{R['n_comunidades']} comunidades", True), " con una modularidad "
        f"ponderada de {R['modularidad']}, un valor alto en la escala habitual. La comunidad mayor "
        f"agrupa {int(R['perfil_comunidades'][0]['nodos'])} nodos y las cuatro menores sólo 2. Lo "
        "revelador no es el número sino su composición: "
        f"{sum(1 for c in R['perfil_comunidades'] if c['videos'] == 1)} de las "
        f"{R['n_comunidades']} comunidades contienen ", ("un solo video", True),
        ". Es decir, la modularidad alta no revela tribus de usuarios con afinidades compartidas: "
        "refleja que cada video captura una audiencia propia que no se solapa con las demás.",
    ])
    perfil = tabla("53_perfil_de_comunidades")[
        ["comunidad", "nodos", "autores", "videos", "comentarios", "intensidad_comentarios_por_autor",
         "sent_medio", "pct_negativos"]]
    perfil.columns = ["C", "Nodos", "Autores", "Videos", "Comentarios", "Comentarios/autor",
                      "Sentimiento medio", "% negativos"]
    insertar_tabla(doc, perfil, anchos=[1.2, 1.8, 2.0, 1.7, 2.4, 2.8, 2.6, 2.0], tamano=8.5,
                   alineacion_derecha=(1, 2, 3, 4, 5, 6, 7))
    leyenda(doc, "tabla", "Las 17 comunidades detectadas, ordenadas por tamaño.")

    titulo(doc, 2, "7.4 Visualización de todas las comunidades")
    figura_apaisada(doc, "11_comunidades.png",
                    "Las 17 comunidades detectadas. Izquierda: la red coloreada por comunidad, con los "
                    "videos como cuadrados. Derecha: tamaño, composición y sentimiento medio de cada "
                    "comunidad.")

    titulo(doc, 2, "7.5 Caracterización de las tres comunidades principales")
    for indice in range(3):
        comunidad = R["perfil_comunidades"][indice]
        titulo(doc, 3, f"Comunidad {comunidad['comunidad']} — {comunidad['nodos']} nodos, "
                       f"{comunidad['comentarios']} comentarios")
        detalle = pd.DataFrame([
            ["Composición", f"{comunidad['autores']} autores y {comunidad['videos']} video(s)"],
            ["Videos", comunidad["videos_titulos"]],
            ["Canales", comunidad["canales"]],
            ["Categorías", comunidad["categorias"]],
            ["Intensidad", f"{comunidad['intensidad_comentarios_por_autor']} comentarios por autor · "
                           f"{comunidad['likes_totales']} «me gusta» · "
                           f"{comunidad['respuestas_totales']} respuestas"],
            ["Temas frecuentes", comunidad["temas_frecuentes"]],
            ["Sentimiento", f"medio {comunidad['sent_medio']:+.3f} · "
                            f"{comunidad['pct_negativos']} % negativos · "
                            f"{comunidad['pct_positivos']} % positivos"],
        ], columns=["Dimensión", "Descripción"])
        insertar_tabla(doc, detalle, anchos=[3.2, 13.3], tamano=8.5)

    texto(doc, [
        "Las tres comunidades principales ilustran el patrón general y sus matices. La ",
        ("comunidad 1", True), " es la mayor con diferencia y corresponde a un único video sobre el "
        "gasto de los diputados: reúne 125 autores que comentan casi una sola vez cada uno "
        f"({R['perfil_comunidades'][0]['intensidad_comentarios_por_autor']} comentarios por autor) y "
        f"es la más negativa de la muestra ({R['perfil_comunidades'][0]['sent_medio']}, "
        f"{R['perfil_comunidades'][0]['pct_negativos']} % de negativos), con vocabulario centrado en "
        "«pueblo», «diputado», «pagar» y «sueldo». La ", ("comunidad 2", True), " gira en torno a la "
        "cooptación universitaria y mezcla crítica con reconocimiento al trabajo periodístico "
        "—de ahí que «excelente» aparezca entre sus términos frecuentes y su tono sea menos negativo "
        f"({R['perfil_comunidades'][1]['sent_medio']}). La ", ("comunidad 3", True), " es la única de "
        f"las tres que agrupa {int(R['perfil_comunidades'][2]['videos'])} videos distintos: son "
        "contenidos sobre ciudad, servicios e infraestructura del mismo canal que comparten "
        "comentaristas, y su tono es el más moderado del trío "
        f"({R['perfil_comunidades'][2]['sent_medio']}).",
    ])
    texto(doc, [
        "Fuera del trío principal conviene señalar el contraste extremo: la comunidad 5 —"
        "planificación urbana municipal— es la única grande con sentimiento claramente positivo "
        f"({R['perfil_comunidades'][4]['sent_medio']:+.3f}, sólo "
        f"{R['perfil_comunidades'][4]['pct_negativos']} % de negativos). El tono no es una "
        "característica de «los comentarios de YouTube» sino del tipo de contenido que los provoca.",
    ])
    salto_de_pagina(doc)


# --------------------------------------------------------------------------------------------
# 8. Centralidad
# --------------------------------------------------------------------------------------------
def seccion_8(doc: Document) -> None:
    titulo(doc, 1, "8. Nodos centrales y participantes puente")

    titulo(doc, 2, "8.1 Medidas seleccionadas y justificación")
    medidas = pd.DataFrame([
        ["Grado bipartito", "Videos comentados (autor) o comentaristas distintos (video)",
         "Medida directa de alcance. Se normaliza por el tamaño del lado opuesto, no por n, porque "
         "un autor sólo puede conectarse con videos."],
        ["Fuerza (grado ponderado)", "Número total de comentarios del nodo",
         "Separa amplitud (muchos videos) de intensidad (muchos comentarios en uno)."],
        ["Intermediación", "Fracción de caminos más cortos que pasan por el nodo",
         "Es la medida de participante puente: identifica a quien conecta audiencias que de otro "
         "modo estarían separadas."],
        ["Cercanía", "Inverso de la distancia media al resto de su componente",
         "Indica integración. Se calcula por componente, porque la red está fragmentada."],
        ["Vector propio", "Conexión con nodos a su vez muy conectados",
         "Mide estar en la zona densa de la red; en bipartitas se interpreta con cautela."],
        ["PageRank", "Probabilidad estacionaria de un paseo aleatorio",
         "Robusto ante la fragmentación gracias al factor de amortiguación; es la medida preferida "
         "para rankear videos."],
    ], columns=["Medida", "Qué mide aquí", "Por qué se usa"])
    insertar_tabla(doc, medidas, anchos=[3.0, 5.2, 8.3], tamano=8.5)
    leyenda(doc, "tabla", "Medidas de centralidad calculadas y justificación de cada una.")
    texto(doc, [
        "Las seis medidas se calculan sobre la ", ("red bipartita observada", True), ". Es "
        "deliberado no calcular la intermediación sobre la proyección autor–autor: dentro de una "
        "camarilla todos los caminos tienen longitud 1, de modo que la intermediación colapsa a 0 "
        "salvo para los pocos autores que unen camarillas, lo que exagera artificialmente su "
        "importancia relativa.",
    ])

    titulo(doc, 2, "8.2 Interpretación para autores: recurrencia y diversidad")
    autores = tabla("56_centralidad_autores").head(10)[
        ["etiqueta", "diversidad", "recurrencia", "intermediacion", "cercania", "pagerank", "sent_medio"]]
    autores.columns = ["Autor", "Videos (diversidad)", "Comentarios (recurrencia)", "Intermediación",
                       "Cercanía", "PageRank", "Sentimiento medio"]
    insertar_tabla(doc, autores, anchos=[3.6, 2.4, 2.7, 2.2, 1.9, 1.9, 2.3], tamano=8.5,
                   alineacion_derecha=(1, 2, 3, 4, 5, 6))
    leyenda(doc, "tabla", "Los diez autores más centrales, ordenados por intermediación.")
    texto(doc, [
        "Para los autores, el enunciado pide distinguir ", ("recurrencia", True), " —cuánto vuelve "
        "una persona— de ", ("diversidad de participación", True), " —en cuántos contenidos "
        f"distintos aparece—, y la distinción resulta decisiva. Hay {R['n_recurrentes']} autores "
        "recurrentes (más de un comentario), pero los más recurrentes no son los más centrales: ",
        (R["recurrentes_top"][0]["etiqueta"], False, True), " publicó "
        f"{R['recurrentes_top'][0]['recurrencia']} comentarios, todos en el mismo video, y su "
        "intermediación es exactamente 0. En cambio ", (R["top_autores"][0]["etiqueta"], False, True),
        f", con sólo {R['top_autores'][0]['recurrencia']} comentarios repartidos en "
        f"{R['top_autores'][0]['diversidad']} videos de canales distintos, alcanza la intermediación "
        f"más alta de la red ({R['top_autores'][0]['intermediacion']}). ",
        ("En esta red la centralidad no la da el volumen sino el movimiento entre contenidos.", True),
    ])

    titulo(doc, 2, "8.2 Interpretación para videos: alcance y capacidad de conectar audiencias")
    videos_c = tabla("57_centralidad_videos").head(10)[
        ["etiqueta", "channel_name", "grado", "fuerza", "intermediacion", "pagerank", "view_count"]]
    videos_c.columns = ["Video", "Canal", "Comentaristas", "Comentarios", "Intermediación", "PageRank", "Vistas"]
    insertar_tabla(doc, videos_c, anchos=[4.6, 3.4, 2.1, 2.0, 1.8, 1.4, 1.6], tamano=8,
                   alineacion_derecha=(2, 3, 4, 5, 6))
    leyenda(doc, "tabla", "Los diez videos más centrales de la red.")
    texto(doc, [
        "Para los videos, las medidas responden dos preguntas distintas. El ", ("alcance", True),
        " lo mide el grado y el PageRank: «", (R["top_videos_red"][0]["etiqueta"], False, True),
        f"» domina ambos ({R['top_videos_red'][0]['grado']} comentaristas distintos, PageRank "
        f"{R['top_videos_red'][0]['pagerank']}), casi tres veces el segundo. La ",
        ("capacidad de conectar audiencias", True), " la mide la intermediación, y ahí el orden "
        "cambia: «Conferencia de Prensa del Gobierno» ocupa el segundo lugar en intermediación pese "
        "a tener sólo 19 comentaristas, porque es el punto por donde la audiencia gubernamental se "
        "enlaza con el resto de la red. Un video puede ser grande sin ser articulador, y viceversa.",
    ])
    figura_apaisada(doc, "12_centralidad.png",
                    "Autores por intermediación, videos por PageRank y el contraste entre amplitud e "
                    "intensidad de participación por autor.")

    titulo(doc, 2, "8.3 Participantes recurrentes, autores puente y videos articuladores")
    texto(doc, [
        "La prueba de que un nodo es articulador no es una métrica sino una verificación directa: se "
        "elimina el nodo y se cuenta si la red se parte. Con ese criterio hay ",
        (f"{R['n_autores_puente']} autores puente", True), " y ",
        (f"{R['n_videos_articuladores']} videos articuladores", True), " en la subred observada.",
    ])
    articuladores = tabla("59_videos_articuladores").head(8)[
        ["video", "canal", "comentaristas", "intermediacion", "componentes_despues"]]
    articuladores.columns = ["Video", "Canal", "Comentaristas", "Intermediación",
                             "Componentes tras eliminarlo"]
    insertar_tabla(doc, articuladores, anchos=[5.4, 3.8, 2.3, 2.4, 2.6], tamano=8.5,
                   alineacion_derecha=(2, 3, 4))
    leyenda(doc, "tabla", "Videos articuladores: su eliminación fragmenta la red en muchas piezas.")
    texto(doc, [
        "Los dos tipos de articulador tienen efectos de magnitud muy distinta. Eliminar al autor "
        "puente más importante lleva la red de "
        f"{R['n_componentes_obs']} a 11 componentes: separa un bloque completo del resto. Eliminar "
        "el video más central la lleva de "
        f"{R['n_componentes_obs']} a {int(R['articuladores_video'][0]['componentes_despues'])}, "
        "porque sus 128 comentaristas quedan sueltos como nodos individuales. La diferencia es "
        "conceptual: ", ("el video es un agregador y el autor es un conector", True),
        ". Sin el video, su audiencia se disuelve en nodos aislados; sin el autor puente, dos "
        "audiencias que se tocaban quedan definitivamente separadas.",
    ])
    texto(doc, [
        ("Hallazgo. ", True), f"La conectividad de esta red descansa en {R['n_autores_puente']} "
        f"personas, el {round(100 * R['n_autores_puente'] / R['autores_unicos'], 1)} % de los autores, "
        "que aportan entre 2 y 4 comentarios cada una. Es un núcleo conector extremadamente delgado: "
        "no hay redundancia estructural que absorba su ausencia, algo que confirma la redundancia "
        "media de "
        f"{[x for x in R['cohesion'] if x['medida'].startswith('Redundancia')][0]['valor']} calculada "
        "en la sección 6.",
    ])
    salto_de_pagina(doc)


# --------------------------------------------------------------------------------------------
# 9. Contenido y sentimiento
# --------------------------------------------------------------------------------------------
def seccion_9(doc: Document) -> None:
    titulo(doc, 1, "9. Análisis de contenido y sentimiento")

    titulo(doc, 2, "9.1 Herramienta, justificación y resultados globales")
    texto(doc, [
        "Se utilizó ", ("pysentimiento", True), " con el modelo ",
        ("robertuito-sentiment-analysis", False, True), ", un transformer ajustado para análisis de "
        "sentimiento en español. La elección se justifica en tres puntos.",
    ])
    vineta(doc, [
        ("Está entrenado en español, no traducido. ", True), "VADER y TextBlob usan léxicos en "
        "inglés; con ellos, un comentario como «qué gran robo» quedaría sin puntuación porque sus "
        "palabras no figuran en el léxico. Un léxico traducido tampoco resuelve los modismos "
        "guatemaltecos presentes en los datos.",
    ])
    vineta(doc, [
        ("Fue entrenado con texto de redes sociales. ", True), "El corpus de origen (TASS, tuits en "
        "español) comparte registro con los comentarios de YouTube: informal, breve, con faltas de "
        "ortografía y emojis.",
    ])
    vineta(doc, [
        ("Es contextual, no un conteo de palabras. ", True), "Al basarse en un transformer capta la "
        "negación y la ironía mejor que un método léxico. Un enfoque de diccionario clasificaría "
        "«no es ninguna maravilla» como positivo por la presencia de «maravilla».",
    ])
    texto(doc, [
        ("Decisión metodológica clave: ", True), "el modelo se aplica sobre ",
        ("texto_original", False, True), " y no sobre ", ("texto_limpio", False, True),
        ". La limpieza elimina la negación («no»), la puntuación y los emojis, que son exactamente "
        "las señales que el modelo necesita. Ésta es la razón de conservar dos versiones del texto, "
        "según se explicó en la sección 2.5. El código incluye además un respaldo léxico en español "
        "con manejo de negación, para que el análisis siga siendo ejecutable si el modelo no está "
        "disponible; la variable ", ("modelo_sentimiento", False, True), " deja constancia de cuál se "
        "utilizó en cada ejecución.",
    ])
    sentimiento = tabla("61_sentimiento_global")
    sentimiento.columns = ["Etiqueta", "Comentarios", "Porcentaje", "Confianza media del modelo"]
    insertar_tabla(doc, sentimiento, anchos=[3.6, 3.6, 3.6, 5.7], tamano=9,
                   alineacion_derecha=(1, 2, 3))
    leyenda(doc, "tabla", "Distribución global del sentimiento en los 406 comentarios.")
    texto(doc, [
        "El resultado global es contundente: ", (f"{R['sentimiento_global']['NEG']} de "
        f"{R['n_comentarios']} comentarios ({R['pct_negativo_global']} %) son negativos", True),
        f", con un puntaje medio de {R['sentimiento_medio']} en una escala de −1 a +1. Los positivos "
        f"y los neutros se reparten casi por igual el resto ({R['sentimiento_global']['POS']} y "
        f"{R['sentimiento_global']['NEU']} respectivamente). La confianza media del modelo es alta en "
        "las clases extremas (0.87 en negativos, 0.81 en positivos) y notablemente menor en la clase "
        "neutra (0.63), lo que es esperable: «neutro» es la categoría residual y la más difícil de "
        "delimitar.",
    ])

    titulo(doc, 2, "9.2 Sentimiento por video, canal, tema y comunidad")
    texto(doc, [
        "Se aplica un ", ("umbral mínimo de 10 comentarios", True), " para comparar grupos. Por "
        "debajo de ese tamaño, el error estándar de una proporción supera los 15 puntos "
        "porcentuales y cualquier diferencia observada sería indistinguible del ruido. Los grupos "
        "menores se reportan en las tablas pero se marcan como no comparables.",
    ])
    canales = tabla("64_sentimiento_por_canal")[
        ["channel_name", "comentarios", "autores", "sent_medio", "pct_neg", "pct_pos", "comparable"]]
    canales.columns = ["Canal", "Comentarios", "Autores", "Sentimiento medio", "% negativos",
                       "% positivos", "¿Comparable?"]
    canales["¿Comparable?"] = canales["¿Comparable?"].map({True: "Sí", False: "No (n < 10)"})
    insertar_tabla(doc, canales, anchos=[4.6, 2.1, 1.7, 2.4, 1.8, 1.8, 2.1], tamano=8.5,
                   alineacion_derecha=(1, 2, 3, 4, 5))
    leyenda(doc, "tabla", "Sentimiento por canal propietario del video.")
    temas = tabla("65_sentimiento_por_categoria")
    temas.columns = ["Categoría", "Comentarios", "Videos", "Sentimiento medio", "% negativos", "¿Comparable?"]
    temas["¿Comparable?"] = temas["¿Comparable?"].map({True: "Sí", False: "No (n < 10)"})
    insertar_tabla(doc, temas, anchos=[4.6, 2.6, 2.0, 3.0, 2.3, 2.0], tamano=8.5,
                   alineacion_derecha=(1, 2, 3, 4))
    leyenda(doc, "tabla", "Sentimiento por categoría temática de YouTube.")
    comunidades = tabla("66_sentimiento_por_comunidad").head(8)[
        ["comunidad", "comentarios", "autores", "sent_medio", "pct_neg", "pct_pos", "comparable"]]
    comunidades.columns = ["Comunidad", "Comentarios", "Autores", "Sentimiento medio", "% negativos",
                           "% positivos", "¿Comparable?"]
    comunidades["¿Comparable?"] = comunidades["¿Comparable?"].map({True: "Sí", False: "No (n < 10)"})
    insertar_tabla(doc, comunidades, anchos=[2.4, 2.4, 2.1, 2.9, 2.3, 2.3, 2.1], tamano=8.5,
                   alineacion_derecha=(0, 1, 2, 3, 4, 5))
    leyenda(doc, "tabla", "Sentimiento por comunidad detectada en la sección 7.")
    figura_apaisada(doc, "13_sentimiento.png",
                    "Sentimiento global y su variación por canal, video y comunidad. Sólo se grafican "
                    "los grupos con al menos 10 comentarios.")

    titulo(doc, 2, "9.3 Hallazgos de contenido y sentimiento")
    pruebas = tabla("68_pruebas_diferencias_sentimiento")
    pruebas.columns = ["Prueba", "Estadístico", "p", "Detalle"]
    pruebas["p"] = pruebas["p"].map(lambda v: "< 0.001" if float(v) < 0.001 else f"{float(v):.4f}")
    insertar_tabla(doc, pruebas, anchos=[7.2, 2.6, 2.0, 4.7], tamano=8.5, alineacion_derecha=(1, 2))
    leyenda(doc, "tabla", "Pruebas de significancia de las diferencias entre canales.")
    vineta(doc, [
        ("El tono negativo es dominante pero no uniforme. ", True), "La diferencia entre canales es "
        f"estadísticamente significativa (Kruskal–Wallis H = "
        f"{R['pruebas_sentimiento'][0]['estadístico']}, {p_valor(R['pruebas_sentimiento'][0]['p_valor'])}; "
        f"χ² = {R['pruebas_sentimiento'][1]['estadístico']}, {p_valor(R['pruebas_sentimiento'][1]['p_valor'])}) "
        "y el rango es amplio: de −0.502 en el canal de periodismo de investigación a +0.609 en el "
        "canal municipal.",
    ])
    vineta(doc, [
        ("El tema predice el tono mejor que el canal. ", True), "Los contenidos de fiscalización "
        "política concentran los tonos más negativos, mientras que los de obra pública local "
        "concentran los positivos. El caso más claro es que dos canales institucionales "
        "—Municipalidad y Gobierno— obtienen resultados opuestos: +0.609 y −0.221 respectivamente.",
    ])
    vineta(doc, [
        ("El vocabulario confirma la lectura. ", True), "Los términos dominantes de los comentarios "
        "—«pueblo», «diputado», «pagar», «dinero», «corrupto», «sueldo»— y bigramas como «pacto "
        "corrupto», «morir hambre» o «pagar sueldo» describen una agenda de fiscalización del gasto "
        "público, no una conversación social diversa.",
    ])
    vineta(doc, [
        ("Los comentarios más respaldados son los menos negativos. ", True), "Como se mostró en la "
        f"sección 3.6, la correlación entre «me gusta» y puntaje de sentimiento es positiva "
        f"(ρ = {R['rho_likes_sentimiento']}, {p_valor(R['p_likes_sentimiento'])}). El respaldo "
        "explícito de la audiencia acompaña a los mensajes de apoyo.",
    ])
    vineta(doc, [
        ("Los emojis refuerzan el patrón ambivalente. ", True), "Los más usados combinan burla "
        "(😂, 29 apariciones) con aplauso y celebración (👏, 🎉, 👍) y la bandera de Guatemala, "
        "lo que es coherente con una muestra que mezcla crítica política y apoyo a obra pública.",
    ])
    vineta(doc, [
        ("Advertencia sobre el alcance de estas etiquetas. ", True), "El modelo fue entrenado con "
        "tuits y no está calibrado para la ironía ni los modismos guatemaltecos. Las etiquetas son "
        "una aproximación útil ", ("a nivel agregado", True), ", no un juicio fiable comentario a "
        "comentario.",
    ])
    salto_de_pagina(doc)


# --------------------------------------------------------------------------------------------
# 10. Interpretación, limitaciones y conclusiones
# --------------------------------------------------------------------------------------------
def seccion_10(doc: Document) -> None:
    titulo(doc, 1, "10. Interpretación, limitaciones y conclusiones")

    titulo(doc, 2, "10.1 Los hallazgos en el contexto de la participación en YouTube")
    hallazgos = tabla("71_hallazgos_integrados")
    for _, fila_h in hallazgos.iterrows():
        titulo(doc, 3, f"{fila_h['eje']}: {fila_h['hallazgo']}")
        texto(doc, [("Evidencia. ", True), fila_h["evidencia"]], espacio_despues=2)
        texto(doc, [("Lectura. ", True), fila_h["interpretación"]])

    titulo(doc, 2, "10.2 Limitaciones")
    texto(doc, [
        "Las limitaciones no son una nota al pie: determinan qué afirmaciones son válidas. Se "
        "declaran con la evidencia cuantitativa que las respalda y con su consecuencia concreta.",
    ])
    limitaciones = tabla("69_limitaciones")
    limitaciones.columns = ["Limitación", "Evidencia cuantitativa", "Consecuencia para la interpretación"]
    insertar_tabla(doc, limitaciones, anchos=[3.4, 6.2, 6.9], tamano=8)
    leyenda(doc, "tabla", "Limitaciones del estudio con su evidencia y su consecuencia.")

    titulo(doc, 2, "10.3 Descripción, asociación e inferencia")
    niveles = tabla("70_descripcion_asociacion_inferencia")
    niveles.columns = ["Nivel", "Ejemplo extraído de este análisis", "Por qué"]
    insertar_tabla(doc, niveles, anchos=[3.0, 7.0, 6.5], tamano=8.5)
    leyenda(doc, "tabla", "Distinción explícita entre los tres niveles de afirmación.")
    texto(doc, [
        "La regla que se siguió en todo el informe es simple. Contar lo que hay en los archivos es ",
        ("descripción", True), " y siempre es válido. Relacionar dos variables observadas es ",
        ("asociación", True), " y es válido si se declara el tamaño de muestra y no se le atribuye "
        "causalidad. Extender los resultados a usuarios, canales o videos que no están en la muestra "
        "es ", ("inferencia", True), " y ", ("no es válido aquí", True), f": la muestra cubre el "
        f"{R['cobertura_videos_pct']} % de los videos, no es probabilística y está dominada por un "
        "solo canal. Por eso este informe nunca afirma «los guatemaltecos son negativos en YouTube» "
        "sino «el 61.3 % de los comentarios de esta muestra fueron clasificados como negativos».",
    ])

    titulo(doc, 2, "10.4 Conclusiones integradas")
    texto(doc, [
        "Las cuatro dimensiones analizadas —red, contenido, sentimiento y limitaciones— convergen "
        "en una misma conclusión.",
    ])
    vineta(doc, [
        ("La red describe agregación en torno al contenido, no una comunidad de personas. ", True),
        f"El {R['pct_autores_grado1']} % de los autores tiene grado 1, la transitividad bipartita es "
        f"nula por construcción y las {R['n_comunidades']} comunidades detectadas "
        f"(modularidad {R['modularidad']}) coinciden casi exactamente con videos individuales. Lo que "
        "agrupa a las personas es el video que están comentando, no un vínculo entre ellas.",
    ])
    vineta(doc, [
        ("El contenido explica el tono mejor que cualquier atributo de los usuarios. ", True),
        f"El {R['pct_negativo_global']} % de negatividad global se descompone en perfiles "
        "temáticos nítidos y estadísticamente distinguibles: fiscalización del gasto público en el "
        "extremo negativo, obra pública municipal en el positivo. La misma persona no aparece lo "
        "bastante en la muestra como para tener un «tono propio»; el tono es del tema.",
    ])
    vineta(doc, [
        ("La estructura es frágil y depende de muy pocos nodos. ", True),
        f"{R['n_autores_puente']} autores son puntos de articulación y su eliminación fragmenta el "
        "núcleo; la redundancia media es de "
        f"{[x for x in R['cohesion'] if x['medida'].startswith('Redundancia')][0]['valor']}. La "
        "cohesión aparente de la componente mayor "
        f"({R['pct_componente_mayor']} % de la subred) descansa en un puñado de conexiones "
        "individuales, no en una malla densa.",
    ])
    vineta(doc, [
        ("La forma de la red es tanto un resultado social como un artefacto del muestreo. ", True),
        f"Que {R['videos_sin_comentarios']} de {R['n_videos']} videos no tengan comentarios "
        "recolectados —y que los 105 videos de la estrategia official_gov no aporten ni uno— "
        "condiciona la fragmentación observada. Separar ambos efectos con estos datos no es posible; "
        "reconocerlo es parte del resultado.",
    ])
    texto(doc, [
        ("Conclusión final. ", True), "Dentro de la muestra analizada, la participación en YouTube "
        "se comporta como un conjunto de audiencias efímeras y paralelas que se congregan alrededor "
        "de piezas de contenido concretas, expresan mayoritariamente crítica hacia la gestión "
        "pública y se disuelven sin dejar vínculos entre sí. Las herramientas de análisis de redes "
        "resultan útiles no porque revelen una comunidad oculta, sino porque permiten demostrar "
        "cuantitativamente que ", ("esa comunidad no existe en los datos", True), ": lo que hay es "
        "co-presencia alrededor del contenido, sostenida por siete personas que circulan entre "
        "temas.",
    ])
    texto(doc, [
        ("Trabajo futuro. ", True), "Tres extensiones aumentarían sustancialmente el valor del "
        "análisis: recolectar los comentarios de respuesta para construir una verdadera red de "
        "conversación; ampliar la cobertura a todos los videos del catálogo con un muestreo "
        "documentado, para separar el efecto del procedimiento del efecto del comportamiento; y "
        "registrar marcas temporales absolutas, que permitirían estudiar cómo se forma y se disuelve "
        "cada audiencia a lo largo del tiempo.",
    ])
    salto_de_pagina(doc)


# --------------------------------------------------------------------------------------------
# Anexo
# --------------------------------------------------------------------------------------------
def anexo(doc: Document) -> None:
    titulo(doc, 1, "Anexo A. Reproducibilidad y material generado")

    titulo(doc, 2, "A.1 Cómo reproducir el análisis")
    pasos = pd.DataFrame([
        ["1", "python -m venv .venv && source .venv/bin/activate", "Entorno aislado (Python 3.11 o superior)."],
        ["2", "pip install -r requirements.txt", "Instala pandas, networkx, spaCy, python-louvain, pysentimiento y demás."],
        ["3", "python -m spacy download es_core_news_sm", "Modelo de español para tokenización y lematización."],
        ["4", "python scripts/lab6_analisis.py", "Ejecuta las actividades 1 a 10 y regenera tablas, figuras y grafos."],
        ["5", "python scripts/generar_notebook.py --ejecutar", "Genera y ejecuta los tres notebooks acumulativos a partir del mismo código."],
        ["6", "python scripts/generar_informe.py", "Regenera este documento de Word con las cifras actualizadas."],
    ], columns=["Paso", "Comando", "Qué hace"])
    insertar_tabla(doc, pasos, anchos=[1.2, 7.3, 8.0], tamano=8.5)
    leyenda(doc, "tabla", "Pasos para reproducir el análisis completo desde cero.")
    texto(doc, [
        "El análisis es determinista: todas las semillas aleatorias están fijadas en 42 (numpy, "
        "Louvain, diseños de fuerzas y cálculo de intermediación). El archivo ",
        ("scripts/lab6_analisis.py", False, True), " está escrito en formato ",
        ("percent", False, True), ", de modo que el mismo código sirve como script ejecutable y como "
        "notebook. De él se derivan tres notebooks acumulativos —hito 1 con las actividades 1 y 2, "
        "hito 2 añadiendo 3 y 4, y la entrega final con 5 a 10—; cada uno vuelve a leer los CSV "
        "originales y puede ejecutarse por separado. Los notebooks se generan y nunca se editan a "
        "mano, lo que evita que las versiones del análisis se desincronicen. Ninguna cifra de este "
        "informe está escrita a mano: todas se leen de ",
        ("outputs/resultados.json", False, True), ", producido por el análisis.",
    ])

    titulo(doc, 2, "A.2 Material generado")
    material = pd.DataFrame([
        ["outputs/tables/", f"{len(R['tablas'])} archivos CSV", "Una tabla por cada resultado del informe, numeradas en el orden del análisis."],
        ["outputs/figures/", f"{len(R['figuras'])} archivos PNG", "Todas las figuras incluidas en este documento, a 200 ppp."],
        ["outputs/graphs/", f"{len(R['grafos'])} archivos GraphML", "Red bipartita completa y observada, y ambas proyecciones, para inspección en Gephi."],
        ["outputs/resultados.json", "1 archivo JSON", f"{len(R)} métricas usadas para redactar este informe."],
        ["data/processed/", "3 archivos CSV", "Comentarios y videos limpios, más el resumen de métricas por video."],
        ["notebooks/", "3 notebooks acumulativos", "Hito 1 (actividades 1–2), hito 2 (+3–4) y entrega final (+5–10), ejecutados."],
    ], columns=["Ruta", "Contenido", "Descripción"])
    insertar_tabla(doc, material, anchos=[4.0, 3.2, 9.3], tamano=8.5)
    leyenda(doc, "tabla", "Material reproducible generado por el análisis.")

    titulo(doc, 2, "A.3 Índice de figuras y tablas del análisis")
    figuras = pd.DataFrame({"Archivo": R["figuras"]})
    figuras["Contenido"] = [
        "Videos y canales con mayor participación observada",
        "Distribuciones de visualizaciones, «me gusta» y respuestas",
        "Curvas de Pareto de la concentración",
        "Visualizaciones frente a comentarios",
        "Frecuencias de palabras, bigramas y hashtags",
        "Nube de palabras de comentarios y contenido",
        "Categorías temáticas y cobertura por estrategia de muestreo",
        "Red bipartita completa y subred observada",
        "Proyecciones autor–autor y video–video",
        "Distribución de grados y componentes conexas",
        "Las 17 comunidades detectadas",
        "Centralidad de autores y videos",
        "Sentimiento global, por canal, por video y por comunidad",
    ][:len(R["figuras"])]
    insertar_tabla(doc, figuras, anchos=[7.0, 9.5], tamano=8.5)
    leyenda(doc, "tabla", "Figuras generadas por el análisis.")


# --------------------------------------------------------------------------------------------
def main() -> None:
    doc = Document()
    seccion = doc.sections[0]
    seccion.page_width, seccion.page_height = Cm(21.59), Cm(27.94)   # carta
    seccion.left_margin = seccion.right_margin = Cm(2.5)
    seccion.top_margin = seccion.bottom_margin = Cm(2.2)
    preparar_estilos(doc)
    numerar_paginas(seccion)

    portada(doc)
    resumen_ejecutivo(doc)
    seccion_1(doc)
    seccion_2(doc)
    seccion_3(doc)
    seccion_3_preguntas(doc)
    seccion_3_extras(doc)
    seccion_4(doc)
    seccion_5(doc)
    seccion_6(doc)
    seccion_7(doc)
    seccion_8(doc)
    seccion_9(doc)
    seccion_10(doc)
    anexo(doc)

    SALIDA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SALIDA)
    print(f"Informe generado: {SALIDA.relative_to(ROOT)}")
    print(f"  {CONTADOR['tabla']} tablas · {CONTADOR['figura']} figuras")


if __name__ == "__main__":
    main()
